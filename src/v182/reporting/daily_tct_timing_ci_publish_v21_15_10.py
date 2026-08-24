from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import json

import pandas as pd

ROOT = Path(__file__).resolve().parents[3]
VERSION = "DAILY_TCT_TIMING_CI_PUBLISH_V21_15_10"
SOURCE_RELATIVE = Path("outputs/daily_tct_ct/TCT_SHADOW_V24_1_7.csv")
CSV_RELATIVE = Path("outputs/committee_master/CI_TCT_T1_T2_CONTEXT.csv")
AUDIT_RELATIVE = Path("outputs/audit/DAILY_TCT_T1_T2_CONTEXT_V21_15_10.json")


def _read(path: Path) -> pd.DataFrame:
    if not path.exists() or path.stat().st_size == 0:
        return pd.DataFrame()
    try:
        return pd.read_csv(path, sep=";", encoding="utf-8-sig", low_memory=False)
    except (pd.errors.EmptyDataError, pd.errors.ParserError, OSError):
        return pd.DataFrame()


def _num(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def build_context(root: Path = ROOT) -> pd.DataFrame:
    frame = _read(root / SOURCE_RELATIVE)
    columns = [
        "family", "name", "isin", "signal_status", "engine_decision", "engine_status",
        "quality_score", "quality_coverage_pct", "quality_threshold", "minimum_coverage_pct",
        "quality_threshold_pass", "coverage_pass", "baseline_rank", "baseline_coverage_pct",
        "baseline_eligible", "setup", "score_influence", "live_execution_allowed",
    ]
    if frame.empty:
        return pd.DataFrame(columns=columns)

    work = frame.copy()
    work["tct_baseline_rank"] = _num(work.get("tct_baseline_rank", pd.Series(index=work.index, dtype=float)))
    work["tct_baseline_coverage"] = _num(work.get("tct_baseline_coverage", pd.Series(index=work.index, dtype=float)))
    work = work[work["tct_baseline_rank"].between(1, 20)].copy()
    if work.empty:
        return pd.DataFrame(columns=columns)

    rows: list[dict] = []
    specs = {
        "T1": ("t1_quality_score", "t1_quality_coverage", 70.0, 80.0, ("T1_STARTER_25_SHADOW", "T1_WATCH_SHADOW")),
        "T2": ("t2_quality_score", "t2_quality_coverage", 75.0, 80.0, ("T2_CONFIRM_75_SHADOW",)),
    }
    for family, (score_col, coverage_col, threshold, min_cov_pct, signal_decisions) in specs.items():
        score = _num(work.get(score_col, pd.Series(index=work.index, dtype=float)))
        coverage = _num(work.get(coverage_col, pd.Series(index=work.index, dtype=float))) * 100.0
        for pos, (_, row) in enumerate(work.iterrows()):
            q = score.iloc[pos]
            cov = coverage.iloc[pos]
            decision = str(row.get("decision") or "")
            is_signal = decision in signal_decisions
            rows.append({
                "family": family,
                "name": str(row.get("name") or ""),
                "isin": str(row.get("isin") or ""),
                "signal_status": "SIGNAL" if is_signal else "PRE_SIGNAL_QUALITY_ONLY",
                "engine_decision": decision,
                "engine_status": str(row.get("status") or ""),
                "quality_score": None if pd.isna(q) else round(float(q), 4),
                "quality_coverage_pct": None if pd.isna(cov) else round(float(cov), 2),
                "quality_threshold": threshold,
                "minimum_coverage_pct": min_cov_pct,
                "quality_threshold_pass": bool(pd.notna(q) and float(q) >= threshold),
                "coverage_pass": bool(pd.notna(cov) and float(cov) >= min_cov_pct),
                "baseline_rank": None if pd.isna(row.get("tct_baseline_rank")) else int(float(row.get("tct_baseline_rank"))),
                "baseline_coverage_pct": None if pd.isna(row.get("tct_baseline_coverage")) else round(float(row.get("tct_baseline_coverage")) * 100.0, 2),
                "baseline_eligible": bool(row.get("baseline_eligible_without_t1_t2", False)),
                "setup": "" if pd.isna(row.get("setup")) else str(row.get("setup")),
                "score_influence": 0.0,
                "live_execution_allowed": False,
            })
    out = pd.DataFrame(rows, columns=columns)
    return out.sort_values(["family", "quality_score", "baseline_rank"], ascending=[True, False, True], na_position="last").reset_index(drop=True)


def _append_word(path: Path, context: pd.DataFrame) -> None:
    if not path.exists():
        return
    from docx import Document
    document = Document(path)
    document.add_page_break()
    document.add_heading("TCT — suivi T1 / T2", level=1)
    document.add_paragraph(
        "Cette section publie le suivi exact du Top20 TCT même lorsqu'aucun signal ne franchit tous les gates. "
        "Les lignes PRE_SIGNAL_QUALITY_ONLY sont des scores diagnostiques de proximité et ne constituent ni un T1, ni un T2, ni une préconisation d'achat. "
        "Influence sur le score et la décision du moteur : 0. Aucun ordre réel."
    )
    for family in ("T1", "T2"):
        subset = context[context["family"].eq(family)].copy()
        signals = int(subset["signal_status"].eq("SIGNAL").sum())
        document.add_heading(f"{family} — {signals} signal(s) confirmé(s)", level=2)
        if subset.empty:
            document.add_paragraph("Aucun titre analysé.")
            continue
        table = document.add_table(rows=1, cols=7)
        headers = ["Titre", "ISIN", "Statut", "Score", "Couverture", "Rang base", "Couv. base"]
        for cell, label in zip(table.rows[0].cells, headers):
            cell.text = label
        for row in subset.itertuples():
            cells = table.add_row().cells
            cells[0].text = str(row.name)
            cells[1].text = str(row.isin)
            cells[2].text = "SIGNAL" if row.signal_status == "SIGNAL" else "pré-signal"
            cells[3].text = "n/a" if pd.isna(row.quality_score) else f"{float(row.quality_score):.2f}"
            cells[4].text = "n/a" if pd.isna(row.quality_coverage_pct) else f"{float(row.quality_coverage_pct):.0f}%"
            cells[5].text = "n/a" if pd.isna(row.baseline_rank) else str(int(row.baseline_rank))
            cells[6].text = "n/a" if pd.isna(row.baseline_coverage_pct) else f"{float(row.baseline_coverage_pct):.2f}%"
    document.save(path)


def _append_excel(path: Path, context: pd.DataFrame) -> None:
    if not path.exists():
        return
    from openpyxl import load_workbook
    book = load_workbook(path)
    for sheet in ("TCT_T1", "TCT_T2"):
        if sheet in book.sheetnames:
            del book[sheet]
    for family, sheet in (("T1", "TCT_T1"), ("T2", "TCT_T2")):
        ws = book.create_sheet(sheet)
        subset = context[context["family"].eq(family)]
        cols = list(context.columns)
        ws.append(cols)
        for values in subset.itertuples(index=False, name=None):
            ws.append(list(values))
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        ws.sheet_view.showGridLines = False
        for col_cells in ws.columns:
            width = min(max(max((len(str(c.value)) if c.value is not None else 0 for c in list(col_cells)[:100]), default=0) + 2, 10), 35)
            ws.column_dimensions[col_cells[0].column_letter].width = width
    book.save(path)


def _append_android(path: Path, context: pd.DataFrame) -> None:
    if not path.exists():
        return
    lines = ["", "---", "", "## TCT — suivi T1 / T2", "", "Contexte uniquement : les pré-signaux ne sont pas des signaux d'achat. Influence score/décision = 0.", ""]
    for family in ("T1", "T2"):
        subset = context[context["family"].eq(family)]
        signals = int(subset["signal_status"].eq("SIGNAL").sum())
        lines.extend([f"### {family} — {signals} signal(s)", ""])
        for row in subset.itertuples():
            score = "n/a" if pd.isna(row.quality_score) else f"{float(row.quality_score):.2f}"
            cov = "n/a" if pd.isna(row.quality_coverage_pct) else f"{float(row.quality_coverage_pct):.0f}%"
            label = "SIGNAL" if row.signal_status == "SIGNAL" else "pré-signal"
            lines.append(f"- **{row.name}** — {label} — score {score} — couverture {cov} — rang baseline {row.baseline_rank}")
        lines.append("")
    with path.open("a", encoding="utf-8") as handle:
        handle.write("\n".join(lines).rstrip() + "\n")


def publish(root: Path = ROOT) -> dict:
    context = build_context(root)
    out = root / CSV_RELATIVE
    out.parent.mkdir(parents=True, exist_ok=True)
    context.to_csv(out, sep=";", index=False, encoding="utf-8-sig")

    word = root / "outputs/committee_master/CI_COMITE_INVESTISSEMENT.docx"
    excel = root / "outputs/committee_master/CI_REFERENTIEL_PONDERE.xlsx"
    android = root / "outputs/mobile/ANDROID_CI_CONTROL_CENTER.md"
    _append_word(word, context)
    _append_excel(excel, context)
    _append_android(android, context)

    t1 = context[context["family"].eq("T1")]
    t2 = context[context["family"].eq("T2")]
    payload = {
        "status": "SUCCESS" if not context.empty else "SUCCESS_EMPTY",
        "version": VERSION,
        "generated_at_utc": datetime.now(timezone.utc).isoformat(),
        "source": str(SOURCE_RELATIVE),
        "rows": int(len(context)),
        "t1_rows": int(len(t1)),
        "t1_signals": int(t1["signal_status"].eq("SIGNAL").sum()) if not t1.empty else 0,
        "t2_rows": int(len(t2)),
        "t2_signals": int(t2["signal_status"].eq("SIGNAL").sum()) if not t2.empty else 0,
        "t1_quality_threshold": 70.0,
        "t2_quality_threshold": 75.0,
        "minimum_quality_coverage_pct": 80.0,
        "score_influence": 0.0,
        "decision_influence": False,
        "t1_t2_scope": "ACTION_TCT_ONLY",
        "real_orders_enabled": False,
        "outputs": {
            "csv": str(CSV_RELATIVE),
            "word": "outputs/committee_master/CI_COMITE_INVESTISSEMENT.docx",
            "excel": "outputs/committee_master/CI_REFERENTIEL_PONDERE.xlsx",
            "android": "outputs/mobile/ANDROID_CI_CONTROL_CENTER.md",
        },
    }
    audit = root / AUDIT_RELATIVE
    audit.parent.mkdir(parents=True, exist_ok=True)
    audit.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    return payload


if __name__ == "__main__":
    print(json.dumps(publish(), ensure_ascii=False, indent=2, default=str))
