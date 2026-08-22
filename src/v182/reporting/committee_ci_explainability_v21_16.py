from __future__ import annotations

from pathlib import Path
import json

import pandas as pd

from v182.decision.committee_master import load_registry
from v182.decision.tct_baseline_v24_1_7 import SETUP_COMPONENT, WEIGHTS_V24_1_2
from v182.decision.tct_baseline_v24_1_8 import ACTIVE_WEIGHT
from v182.decision.tct_timing_exact_v24_1_7 import T1_WEIGHTS, T2_WEIGHTS
from v182.reporting import committee_ci_explainability as legacy

ROOT = Path(__file__).resolve().parents[3]
CI_VERSION = "CI_RESTITUTION_V21_16_3"
TCT_CI_CODES = {"T1_STARTER_25_SHADOW", "T1_WATCH_SHADOW", "T2_CONFIRM_75_SHADOW"}
CI_SELECTED_CODES = set(legacy.SELECTED_CODES) | TCT_CI_CODES
INTERNAL_PROVENANCE_SOURCES = {"TCT_V24_1_7_EXACT_COMPONENTS", "TCT_BASELINE_V24_1_8"}
TCT_MIN_EXACT_COMPONENT_WEIGHT = 0.80


def _bool(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "oui"}


def _num(value) -> float | None:
    parsed = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
    return None if pd.isna(parsed) else float(parsed)


def _tct_family(decision: str, setup: object) -> tuple[str, dict[str, float], str] | None:
    decision = str(decision or "").upper()
    setup_text = str(setup or "").upper()
    if decision in {"T1_STARTER_25_SHADOW", "T1_WATCH_SHADOW"} or setup_text == "T1":
        return "T1", T1_WEIGHTS, "t1_component_"
    if decision == "T2_CONFIRM_75_SHADOW" or setup_text in {"T2", "T2_CONFIRMATION", "T2_EXACT_TIMING_CONFIRMATION"}:
        return "T2", T2_WEIGHTS, "t2_component_"
    return None


def _tct_exact_details(selected: pd.DataFrame) -> pd.DataFrame:
    chosen = selected[
        selected["asset_class"].astype(str).eq("ACTION")
        & selected["horizon"].astype(str).eq("TCT")
        & selected["decision"].astype(str).isin(TCT_CI_CODES)
    ]
    if chosen.empty:
        return pd.DataFrame()
    rows: list[dict] = []
    for _, decision in chosen.iterrows():
        family = _tct_family(str(decision.get("decision") or ""), decision.get("setup"))
        if family is None:
            continue
        label, weights, prefix = family
        values = {name: _num(decision.get(prefix + name)) for name in weights}
        observed_weight = sum(float(weight) for name, weight in weights.items() if values[name] is not None)
        for name, weight in weights.items():
            value = values[name]
            available = value is not None and observed_weight > 0
            effective = float(weight) / observed_weight if available else 0.0
            rows.append(
                {
                    "asset_class": "ACTION",
                    "horizon": "TCT",
                    "isin": str(decision.get("isin") or ""),
                    "name": decision.get("name"),
                    "decision": decision.get("decision"),
                    "final_score": decision.get("score"),
                    "criterion": f"TCT_{label}_{name.upper()}",
                    "criterion_status": "ACTIVE" if available else "MISSING",
                    "raw_value": value,
                    "direction": "HIGHER_BETTER",
                    "resolution": f"DERIVED:{prefix}{name}",
                    "source_field": prefix + name,
                    "criterion_score_0_100": value,
                    "theoretical_weight_pct": float(weight) * 100.0,
                    "effective_weight_pct": effective * 100.0,
                    "weighted_contribution_points": value * effective if available else None,
                    "contribution_scope": "FINAL_SCORE",
                    "source": "TCT_V24_1_7_EXACT_COMPONENTS",
                    "source_url": None,
                    "as_of": decision.get("generated_at_utc"),
                    "evidence_level": "A_INTERNAL_GOVERNED",
                    "validation_status": f"{label}_EXACT_COMPONENTS_AVAILABLE_WEIGHT_RENORMALIZED",
                }
            )
    return pd.DataFrame(rows)


def _tct_baseline_context(root: Path, selected: pd.DataFrame) -> pd.DataFrame:
    chosen = selected[
        selected["asset_class"].astype(str).eq("ACTION")
        & selected["horizon"].astype(str).eq("TCT")
        & selected["decision"].astype(str).isin(TCT_CI_CODES)
    ]
    baseline = legacy._read(root / "outputs" / "committee_master" / "TCT_BASELINE_V24_1_8.csv")
    if chosen.empty or baseline.empty or "isin" not in baseline:
        return pd.DataFrame()
    baseline = baseline.drop_duplicates("isin", keep="last").copy()
    baseline["isin"] = baseline["isin"].astype(str)
    lookup = baseline.set_index("isin", drop=False)
    active_weights = {name: float(weight) for name, weight in WEIGHTS_V24_1_2.items() if name != SETUP_COMPONENT}
    rows: list[dict] = []
    for _, decision in chosen.iterrows():
        isin = str(decision.get("isin") or "")
        if isin not in lookup.index:
            continue
        source_row = lookup.loc[isin]
        if isinstance(source_row, pd.DataFrame):
            source_row = source_row.iloc[-1]
        for name, weight in active_weights.items():
            value = _num(source_row.get(f"tct_baseline_component_{name}"))
            rows.append(
                {
                    "asset_class": "ACTION",
                    "horizon": "TCT",
                    "isin": isin,
                    "name": decision.get("name"),
                    "decision": decision.get("decision"),
                    "final_score": decision.get("score"),
                    "criterion": f"TCT_BASELINE_GATE_{name.upper()}",
                    "criterion_status": "CONTEXT_GATE" if value is not None else "CONTEXT_MISSING",
                    "raw_value": value,
                    "direction": "HIGHER_BETTER",
                    "resolution": f"DERIVED:tct_baseline_component_{name}",
                    "source_field": f"tct_baseline_component_{name}",
                    "criterion_score_0_100": value,
                    "theoretical_weight_pct": (weight / ACTIVE_WEIGHT * 100.0) if ACTIVE_WEIGHT else 0.0,
                    "effective_weight_pct": 0.0,
                    "weighted_contribution_points": None,
                    "contribution_scope": "CONTEXT_GATE_NOT_FINAL_SCORE",
                    "source": "TCT_BASELINE_V24_1_8",
                    "source_url": None,
                    "as_of": decision.get("generated_at_utc"),
                    "evidence_level": "A_INTERNAL_GOVERNED",
                    "validation_status": "BASELINE_TOP20_AND_COVERAGE_PREREQUISITE_CONTEXT_ONLY",
                }
            )
    return pd.DataFrame(rows)


def _attach_provenance_preserving_internal(root: Path, detail: pd.DataFrame) -> pd.DataFrame:
    if detail.empty:
        return detail
    ordered = detail.copy()
    ordered["_ci_row_order"] = range(len(ordered))
    internal_mask = ordered.get("source", pd.Series("", index=ordered.index)).astype(str).isin(INTERNAL_PROVENANCE_SOURCES)
    internal = ordered[internal_mask].copy()
    external = ordered[~internal_mask].copy()
    if not external.empty:
        external = legacy._attach_provenance(root, external)
    merged = pd.concat([external, internal], ignore_index=True, sort=False)
    return merged.sort_values("_ci_row_order", kind="stable").drop(columns=["_ci_row_order"]).reset_index(drop=True)


def _tct_reference_integrity(selected: pd.DataFrame, detail: pd.DataFrame) -> dict:
    tct_selected = selected[
        selected["asset_class"].astype(str).eq("ACTION")
        & selected["horizon"].astype(str).eq("TCT")
        & selected["decision"].astype(str).isin(TCT_CI_CODES)
    ]
    tct_keys = legacy._selection_keys(tct_selected)
    if not tct_keys:
        return {
            "selected_tct_keys": 0,
            "missing_exact_keys": [],
            "missing_baseline_keys": [],
            "undercovered_exact_keys": [],
            "complete": True,
        }
    source = detail.get("source", pd.Series("", index=detail.index)).astype(str) if not detail.empty else pd.Series(dtype=str)
    exact = detail[source.eq("TCT_V24_1_7_EXACT_COMPONENTS")].copy() if not detail.empty else pd.DataFrame()
    baseline = detail[source.eq("TCT_BASELINE_V24_1_8")].copy() if not detail.empty else pd.DataFrame()
    exact_keys = legacy._selection_keys(exact)
    baseline_keys = legacy._selection_keys(baseline)
    missing_exact = sorted(tct_keys - exact_keys)
    missing_baseline = sorted(tct_keys - baseline_keys)
    undercovered: list[tuple] = []
    for key in sorted(tct_keys & exact_keys):
        asset, horizon, isin = key
        subset = exact[
            exact["asset_class"].astype(str).eq(asset)
            & exact["horizon"].astype(str).eq(horizon)
            & exact["isin"].astype(str).eq(isin)
        ]
        active = subset[subset["criterion_status"].astype(str).eq("ACTIVE")]
        theoretical = pd.to_numeric(active.get("theoretical_weight_pct"), errors="coerce").fillna(0.0).sum() / 100.0
        if theoretical + 1e-12 < TCT_MIN_EXACT_COMPONENT_WEIGHT:
            undercovered.append(key)
    return {
        "selected_tct_keys": int(len(tct_keys)),
        "missing_exact_keys": [list(key) for key in missing_exact],
        "missing_baseline_keys": [list(key) for key in missing_baseline],
        "undercovered_exact_keys": [list(key) for key in undercovered],
        "minimum_exact_component_weight": TCT_MIN_EXACT_COMPONENT_WEIGHT,
        "complete": not missing_exact and not missing_baseline and not undercovered,
    }


def _source_validation(root: Path, selected: pd.DataFrame) -> pd.DataFrame:
    if selected.empty:
        return pd.DataFrame()
    frame = selected.copy()
    entry = legacy._read(root / "outputs" / "committee_master" / "V21_8_ENTRY_EXIT_CHALLENGER.csv")
    keys = ["asset_class", "horizon", "isin"]
    if not entry.empty and all(key in entry for key in keys):
        keep = keys + [c for c in ("v21_8_entry_state", "v21_8_entry_reasons", "setup", "t1_source_event_id") if c in entry]
        frame = frame.merge(entry[keep].drop_duplicates(keys, keep="last"), on=keys, how="left", suffixes=("", "_entry"), sort=False, validate="many_to_one")

    decision = frame["decision"].astype(str)
    source_buy_ok = frame.get("ci_source_eligible", pd.Series(False, index=frame.index)).map(_bool)
    source_full = frame.get("source_fully_validated", pd.Series(False, index=frame.index)).map(_bool)
    frame["ci_final_status"] = "SURVEILLANCE_INTERNE"
    frame.loc[decision.eq("BUY_CANDIDATE") & ~source_buy_ok, "ci_final_status"] = "BUY_INTERNE_ATTENTE_SOURCES"
    frame.loc[decision.eq("BUY_CANDIDATE") & source_buy_ok, "ci_final_status"] = "RECOMMANDATION_TOTALEMENT_VALIDEE"
    frame.loc[decision.isin({"T1_STARTER_25_SHADOW", "T1_WATCH_SHADOW"}), "ci_final_status"] = "TCT_T1_SURVEILLANCE"
    frame.loc[decision.eq("T2_CONFIRM_75_SHADOW") & ~source_full, "ci_final_status"] = "TCT_T2_ATTENTE_SOURCES"
    frame.loc[decision.eq("T2_CONFIRM_75_SHADOW") & source_full, "ci_final_status"] = "TCT_T2_SOURCE_CONFIRMED"

    columns = [
        "asset_class", "horizon", "name", "isin", "decision", "score", "coverage_pct", "ci_final_status",
        "source_validation_state", "source_validation_reasons", "source_fully_validated", "ci_source_eligible",
        "boursorama_priority_ready", "boursorama_context_coverage_pct", "boursorama_dynamic_age_hours",
        "boursorama_performance_age_hours", "boursorama_deep_age_hours", "boursorama_etf_dynamic_age_hours",
        "boursorama_etf_deep_age_hours", "boursorama_consensus", "boursorama_n_analysts",
        "boursorama_target_median", "boursorama_target_upside_pct", "boursorama_consensus_delta_7d",
        "boursorama_consensus_delta_1m", "boursorama_target_price_delta_7d_pct",
        "boursorama_target_price_delta_1m_pct", "boursorama_buy_ratio_delta_7d",
        "boursorama_buy_ratio_delta_1m", "boursorama_perf_1d_pct", "boursorama_perf_1w_pct",
        "boursorama_perf_1m_pct", "boursorama_perf_6m_pct", "boursorama_perf_1y_pct",
        "boursorama_estimated_per", "boursorama_estimated_yield_pct", "boursorama_eps_est_2026",
        "boursorama_eps_est_2027", "boursorama_revenue_m_est_2026", "boursorama_revenue_m_est_2027",
        "boursorama_etf_aum_eur_m", "boursorama_etf_morningstar_category", "boursorama_etf_replication",
        "boursorama_etf_management_fee_pct", "boursorama_etf_sri", "boursorama_etf_volatility_1y_pct",
        "boursorama_etf_beta_1y", "boursorama_source_urls", "boursorama_latest_collected_at",
        "investing_required_timeframe", "investing_required_state", "investing_horizon_signal",
        "investing_timing_confirmed", "investing_daily_signal", "investing_weekly_signal",
        "investing_monthly_signal", "investing_age_hours", "investing_source_urls",
        "investing_latest_collected_at", "v21_8_entry_state", "v21_8_entry_reasons", "setup", "t1_source_event_id",
    ]
    return frame[[c for c in columns if c in frame]].copy()


def _report_context(context: pd.DataFrame, source: pd.DataFrame) -> pd.DataFrame:
    """Translate only human-facing labels; canonical Committee decisions remain immutable."""
    if context.empty or source.empty:
        return context
    keys = ["asset_class", "horizon", "isin"]
    status = source[keys + ["ci_final_status"]].drop_duplicates(keys, keep="last")
    before = len(context)
    out = context.merge(status, on=keys, how="left", sort=False, validate="many_to_one")
    if len(out) != before:
        raise RuntimeError("CI_REPORT_CONTEXT_ROW_COUNT_MUTATION")
    pending = out["ci_final_status"].astype(str).eq("BUY_INTERNE_ATTENTE_SOURCES") & out["decision"].astype(str).eq("BUY_CANDIDATE")
    out.loc[pending, "decision"] = "BUY_WAIT_SOURCE_CONFIRMATION"
    return out


def _append_excel_source_sheets(path: Path, source: pd.DataFrame) -> list[str]:
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, Font
    from openpyxl.utils import get_column_letter

    workbook = load_workbook(path)
    sheet_names = ("Validation_sources", "Recommandations_validees", "Attente_sources", "TCT_signaux", "Surveillance")
    for name in sheet_names:
        if name in workbook.sheetnames:
            del workbook[name]
    status = source.get("ci_final_status", pd.Series(dtype=str)).astype(str) if not source.empty else pd.Series(dtype=str)
    groups = [
        ("Validation_sources", source),
        ("Recommandations_validees", source[status.eq("RECOMMANDATION_TOTALEMENT_VALIDEE")] if not source.empty else source),
        ("Attente_sources", source[status.isin({"BUY_INTERNE_ATTENTE_SOURCES", "TCT_T2_ATTENTE_SOURCES"})] if not source.empty else source),
        ("TCT_signaux", source[status.isin({"TCT_T1_SURVEILLANCE", "TCT_T2_ATTENTE_SOURCES", "TCT_T2_SOURCE_CONFIRMED"})] if not source.empty else source),
        ("Surveillance", source[status.eq("SURVEILLANCE_INTERNE")] if not source.empty else source),
    ]
    for sheet_name, frame in groups:
        sheet = workbook.create_sheet(sheet_name)
        if frame.empty:
            sheet.cell(1, 1, "Aucune ligne")
            continue
        for col_idx, column in enumerate(frame.columns, start=1):
            cell = sheet.cell(1, col_idx, str(column))
            cell.font = Font(bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        for row_idx, values in enumerate(frame.itertuples(index=False, name=None), start=2):
            for col_idx, value in enumerate(values, start=1):
                sheet.cell(row_idx, col_idx, None if pd.isna(value) else value)
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for idx, column in enumerate(frame.columns, start=1):
            sample = [len(str(column))] + [len(str(v)) for v in frame[column].dropna().astype(str).head(80)]
            sheet.column_dimensions[get_column_letter(idx)].width = min(max(sample) + 2, 46)
    workbook.save(path)
    return ["Referentiel_pondere", *sheet_names]


def _append_word_source_section(path: Path, source: pd.DataFrame) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK

    document = Document(path)
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("Validation finale Boursorama + Investing — V21.16", level=1)
    document.add_paragraph(
        "Une recommandation d'achat totalement validée exige BUY_CANDIDATE + contexte Boursorama frais et complet + "
        "Investing STRONG_BUY sur l'horizon requis. Pour TCT, T1 reste un signal de surveillance et T2 reste un signal "
        "de décision-support SHADOW : même source-confirmé, il n'est pas requalifié en BUY de production."
    )
    if source.empty:
        document.add_paragraph("Aucune présélection Action/ETF à documenter.")
        document.save(path)
        return

    status = source["ci_final_status"].astype(str)
    groups = [
        ("Recommandations totalement validées", source[status.eq("RECOMMANDATION_TOTALEMENT_VALIDEE")]),
        ("BUY internes en attente de confirmation", source[status.eq("BUY_INTERNE_ATTENTE_SOURCES")]),
        ("TCT T2 source-confirmé — décision-support SHADOW", source[status.eq("TCT_T2_SOURCE_CONFIRMED")]),
        ("TCT T2 en attente de sources", source[status.eq("TCT_T2_ATTENTE_SOURCES")]),
        ("TCT T1 — surveillance", source[status.eq("TCT_T1_SURVEILLANCE")]),
        ("Surveillance / review", source[status.eq("SURVEILLANCE_INTERNE")]),
    ]
    document.add_paragraph(" | ".join(f"{title}: {len(frame)}" for title, frame in groups))

    for title, frame in groups:
        document.add_heading(title, level=2)
        if frame.empty:
            document.add_paragraph("Aucune ligne.")
            continue
        for _, row in frame.iterrows():
            name = str(row.get("name") or row.get("isin") or "N/A")
            score = _num(row.get("score"))
            document.add_heading(f"{row.get('asset_class', 'N/A')} — {name} — {row.get('horizon', 'N/A')}", level=3)
            document.add_paragraph(
                f"Décision interne : {row.get('decision', 'N/A')} | score {'N/A' if score is None else f'{score:.1f}'} | "
                f"statut final CI : {row.get('ci_final_status', 'N/A')}"
            )
            b_ready = _bool(row.get("boursorama_priority_ready"))
            b_cov = _num(row.get("boursorama_context_coverage_pct"))
            document.add_paragraph(
                f"Boursorama : {'VALIDÉ' if b_ready else 'INCOMPLET'} | couverture contexte "
                f"{'N/A' if b_cov is None else f'{b_cov:.1f}%'} | as-of {row.get('boursorama_latest_collected_at', 'N/A')}"
            )
            document.add_paragraph(
                f"Investing requis : {row.get('investing_required_timeframe', 'N/A')}=STRONG_BUY | "
                f"observé : {row.get('investing_horizon_signal', 'N/A')} | Daily={row.get('investing_daily_signal', 'N/A')} | "
                f"Weekly={row.get('investing_weekly_signal', 'N/A')} | Monthly={row.get('investing_monthly_signal', 'N/A')} | "
                f"as-of {row.get('investing_latest_collected_at', 'N/A')}"
            )
            if pd.notna(row.get("v21_8_entry_state")):
                document.add_paragraph(f"Entrée V21.8 : {row.get('v21_8_entry_state')} | {row.get('v21_8_entry_reasons', '')}")
            if pd.notna(row.get("source_validation_reasons")) and str(row.get("source_validation_reasons")) != "OK":
                document.add_paragraph(f"Points restant à confirmer : {row.get('source_validation_reasons')}")
            if pd.notna(row.get("boursorama_source_urls")):
                document.add_paragraph(f"Provenance Boursorama : {row.get('boursorama_source_urls')}")
            if pd.notna(row.get("investing_source_urls")):
                document.add_paragraph(f"Provenance Investing : {row.get('investing_source_urls')}")
    document.save(path)


def _append_android_source_section(path: Path, source: pd.DataFrame) -> None:
    lines = ["", "## Validation finale Boursorama + Investing", ""]
    if source.empty:
        lines.append("Aucune présélection à qualifier.")
    else:
        priority = {
            "RECOMMANDATION_TOTALEMENT_VALIDEE": 0,
            "TCT_T2_SOURCE_CONFIRMED": 1,
            "BUY_INTERNE_ATTENTE_SOURCES": 2,
            "TCT_T2_ATTENTE_SOURCES": 3,
            "TCT_T1_SURVEILLANCE": 4,
            "SURVEILLANCE_INTERNE": 5,
        }
        ordered = source.copy()
        ordered["_rank"] = ordered["ci_final_status"].map(priority).fillna(9)
        ordered["_score"] = pd.to_numeric(ordered.get("score"), errors="coerce")
        ordered = ordered.sort_values(["_rank", "_score"], ascending=[True, False])
        for _, row in ordered.iterrows():
            lines.append(
                f"- **{row.get('asset_class')} {row.get('name') or row.get('isin')} [{row.get('horizon')}]** — "
                f"{row.get('ci_final_status')} — gate {row.get('source_validation_state')} — "
                f"Investing {row.get('investing_required_timeframe')}={row.get('investing_horizon_signal')} — "
                f"entrée {row.get('v21_8_entry_state', 'N/A')}"
            )
    with path.open("a", encoding="utf-8") as handle:
        handle.write("\n".join(lines) + "\n")


def run(root: Path = ROOT) -> dict:
    decisions = legacy._read(root / "outputs" / "committee_master" / "COMMITTEE_DECISIONS.csv")
    if decisions.empty:
        return {"status": "BLOCKED_COMMITTEE_DECISIONS_MISSING", "real_orders_enabled": False}
    selected = decisions[decisions["decision"].astype(str).isin(CI_SELECTED_CODES)].copy()
    selected = selected[selected["asset_class"].astype(str).isin({"ACTION", "ETF"})]

    action_source = legacy._read(root / "outputs" / "V18.2_PEA_ACTIONS_MASTER_ENRICHED.csv")
    etf_source = legacy._read(root / "outputs" / "V18.2_PEA_ETF_MASTER_ENRICHED.csv")
    action_registry = load_registry(root / "config" / "V21_ACTIONS_REFERENCE_V21_0.json")
    etf_registry = load_registry(root / "config" / "V20_7_1_ETF_CRITERIA_REGISTRY.json")
    parts = [
        legacy._generic_details(action_source, selected, action_registry, "ACTION", ["CT", "MT", "SHORT", "TOP_DOWN"]),
        _tct_exact_details(selected),
        _tct_baseline_context(root, selected),
        legacy._generic_details(etf_source, selected, etf_registry, "ETF", ["CT", "SHORT", "TOP_DOWN"]),
        legacy._etf_mt_details(root, selected),
    ]
    detail = pd.concat([part for part in parts if not part.empty], ignore_index=True) if any(not part.empty for part in parts) else pd.DataFrame()
    detail = _attach_provenance_preserving_internal(root, detail)
    context = legacy._join_context(root, selected)

    mobile_dir = root / "outputs" / "mobile"
    committee_dir = root / "outputs" / "committee_master"
    audit_dir = root / "outputs" / "audit"
    mobile_dir.mkdir(parents=True, exist_ok=True)
    committee_dir.mkdir(parents=True, exist_ok=True)
    audit_dir.mkdir(parents=True, exist_ok=True)

    selected_keys = legacy._selection_keys(selected)
    detail_keys = legacy._selection_keys(detail)
    missing_reference = sorted(selected_keys - detail_keys)
    tct_integrity = _tct_reference_integrity(selected, detail)
    if missing_reference or not tct_integrity["complete"]:
        blocked = {
            "status": "BLOCKED_CI_TCT_REFERENCE_INCOMPLETE" if not tct_integrity["complete"] else "BLOCKED_CI_REFERENCE_INCOMPLETE",
            "version": CI_VERSION,
            "selected_rows": int(len(selected)),
            "missing_reference_keys": [list(key) for key in missing_reference],
            "tct_reference_integrity": tct_integrity,
            "score_or_decision_mutation": False,
            "weight_or_threshold_changes": False,
            "external_collection_calls": 0,
            "real_orders_enabled": False,
        }
        (audit_dir / "CI_EXPLAINABILITY_AUDIT.json").write_text(json.dumps(blocked, ensure_ascii=False, indent=2), encoding="utf-8")
        raise RuntimeError(
            "CI_TCT_REFERENCE_INCOMPLETE:"
            f"generic={missing_reference[:10]} exact={tct_integrity['missing_exact_keys'][:10]} "
            f"baseline={tct_integrity['missing_baseline_keys'][:10]} undercovered={tct_integrity['undercovered_exact_keys'][:10]}"
        )

    source = _source_validation(root, selected)
    report_context = _report_context(context, source)
    android_path = mobile_dir / "ANDROID_CI_CONTROL_CENTER.md"
    word_path = committee_dir / "CI_COMITE_INVESTISSEMENT.docx"
    excel_path = committee_dir / "CI_REFERENTIEL_PONDERE.xlsx"
    android_path.write_text(legacy._android(report_context, detail), encoding="utf-8")
    legacy._write_word_report(word_path, report_context, detail)
    legacy._write_excel_reference(excel_path, detail)
    sheets = _append_excel_source_sheets(excel_path, source)
    _append_word_source_section(word_path, source)
    _append_android_source_section(android_path, source)

    reconstruction = legacy._reconstruction(selected, detail)
    status = source.get("ci_final_status", pd.Series(dtype=str)).astype(str) if not source.empty else pd.Series(dtype=str)
    metrics = {
        "fully_validated_recommendations": int(status.eq("RECOMMANDATION_TOTALEMENT_VALIDEE").sum()),
        "internal_buy_waiting_sources": int(status.eq("BUY_INTERNE_ATTENTE_SOURCES").sum()),
        "tct_t1_surveillance": int(status.eq("TCT_T1_SURVEILLANCE").sum()),
        "tct_t2_waiting_sources": int(status.eq("TCT_T2_ATTENTE_SOURCES").sum()),
        "tct_t2_source_confirmed": int(status.eq("TCT_T2_SOURCE_CONFIRMED").sum()),
        "surveillance_rows": int(status.eq("SURVEILLANCE_INTERNE").sum()),
    }
    source.to_csv(committee_dir / "CI_VALIDATION_SOURCES.csv", sep=";", index=False, encoding="utf-8-sig")
    payload = {
        "status": "SUCCESS",
        "version": CI_VERSION,
        "selected_rows": int(len(selected)),
        **metrics,
        "criteria_detail_rows": int(len(detail)),
        "tct_reference_integrity": tct_integrity,
        "tct_exact_reference_included": bool((detail.get("source", pd.Series(dtype=str)).astype(str) == "TCT_V24_1_7_EXACT_COMPONENTS").any()),
        "tct_baseline_context_included": bool((detail.get("source", pd.Series(dtype=str)).astype(str) == "TCT_BASELINE_V24_1_8").any()),
        "tct_baseline_contributes_to_timing_score": False,
        "tct_internal_provenance_preserved": True,
        "android_output": str(android_path.relative_to(root)),
        "word_output": str(word_path.relative_to(root)),
        "excel_output": str(excel_path.relative_to(root)),
        "source_validation_output": "outputs/committee_master/CI_VALIDATION_SOURCES.csv",
        "same_canonical_run_android_word_excel": True,
        "same_selected_set_word_excel": selected_keys == detail_keys,
        "reference_complete_for_selected": selected_keys <= detail_keys and tct_integrity["complete"],
        "excel_visible_sheets": sheets,
        "score_or_decision_mutation": False,
        "weight_or_threshold_changes": False,
        "display_decision_translation_only": True,
        "external_collection_calls": 0,
        "source_validation_read_only": True,
        "t1_t2_scope": "ACTION_TCT_ONLY",
        "tct_t2_remains_shadow_decision_support": True,
        "real_orders_enabled": False,
        "reconstruction": reconstruction,
    }
    (audit_dir / "CI_EXPLAINABILITY_AUDIT.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


if __name__ == "__main__":
    print(json.dumps(run(), ensure_ascii=False, indent=2))
