from __future__ import annotations

from pathlib import Path
import argparse
import json

import pandas as pd

from v182.reporting import decision_brief as legacy

ROOT = Path(__file__).resolve().parents[3]
VERSION = "CI_DECISION_BRIEF_V21_16_3"
TCT_EXACT_CODES = {"T1_STARTER_25_SHADOW", "T1_WATCH_SHADOW", "T2_CONFIRM_75_SHADOW"}
BRIEF_SELECTED_CODES = set(legacy.SELECTED_DECISIONS) | TCT_EXACT_CODES
KEYS = ["asset_class", "horizon", "isin"]
SOURCE_COLUMNS = [
    "asset_class", "horizon", "isin", "source_validation_state", "source_validation_reasons",
    "source_fully_validated", "ci_source_eligible", "boursorama_priority_ready", "boursorama_context_coverage_pct",
    "boursorama_latest_collected_at", "investing_required_timeframe", "investing_horizon_signal",
    "investing_daily_signal", "investing_weekly_signal", "investing_monthly_signal", "investing_age_hours",
    "investing_latest_collected_at",
]


def _read(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, sep=";", encoding="utf-8-sig", low_memory=False)


def _bool(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "oui"}


def _num(value) -> float | None:
    parsed = pd.to_numeric(pd.Series([value]), errors="coerce").iloc[0]
    return None if pd.isna(parsed) else float(parsed)


def _source_rows(root: Path) -> pd.DataFrame:
    decisions = _read(root / "outputs" / "committee_master" / "COMMITTEE_DECISIONS.csv")
    if decisions.empty:
        return pd.DataFrame()
    selected = decisions[decisions["decision"].astype(str).isin(BRIEF_SELECTED_CODES)].copy()
    selected = selected[selected["asset_class"].astype(str).isin({"ACTION", "ETF"})]
    if selected.empty:
        return selected
    decision = selected["decision"].astype(str)
    source_buy_ok = selected.get("ci_source_eligible", pd.Series(False, index=selected.index)).map(_bool)
    source_full = selected.get("source_fully_validated", pd.Series(False, index=selected.index)).map(_bool)
    selected["ci_final_status"] = "SURVEILLANCE_INTERNE"
    selected.loc[decision.eq("BUY_CANDIDATE") & ~source_buy_ok, "ci_final_status"] = "BUY_INTERNE_ATTENTE_SOURCES"
    selected.loc[decision.eq("BUY_CANDIDATE") & source_buy_ok, "ci_final_status"] = "RECOMMANDATION_TOTALEMENT_VALIDEE"
    selected.loc[decision.isin({"T1_STARTER_25_SHADOW", "T1_WATCH_SHADOW"}), "ci_final_status"] = "TCT_T1_SURVEILLANCE"
    selected.loc[decision.eq("T2_CONFIRM_75_SHADOW") & ~source_full, "ci_final_status"] = "TCT_T2_ATTENTE_SOURCES"
    selected.loc[decision.eq("T2_CONFIRM_75_SHADOW") & source_full, "ci_final_status"] = "TCT_T2_SOURCE_CONFIRMED"
    return selected


def _tct_prior(snapshot: pd.DataFrame) -> pd.DataFrame:
    if snapshot.empty or "decision" not in snapshot:
        return pd.DataFrame()
    mask = snapshot["decision"].astype(str).isin(TCT_EXACT_CODES)
    if "horizon" in snapshot:
        mask &= snapshot["horizon"].astype(str).str.upper().eq("TCT")
    return snapshot[mask].copy().drop_duplicates(KEYS, keep="last") if all(key in snapshot for key in KEYS) else pd.DataFrame()


def _legacy_snapshot_only(snapshot: pd.DataFrame) -> pd.DataFrame:
    if snapshot.empty or "decision" not in snapshot:
        return snapshot.copy()
    return snapshot[~snapshot["decision"].astype(str).isin(TCT_EXACT_CODES)].copy()


def _tct_history(source: pd.DataFrame, prior_tct: pd.DataFrame) -> pd.DataFrame:
    out = source.copy()
    for column in ("previous_decision", "previous_score", "score_delta", "tct_change_state"):
        if column not in out:
            out[column] = None
    if out.empty:
        return out
    prior_lookup = {}
    if not prior_tct.empty and all(key in prior_tct for key in KEYS):
        prior_lookup = {
            tuple(str(row.get(key) or "") for key in KEYS): row.to_dict()
            for _, row in prior_tct.drop_duplicates(KEYS, keep="last").iterrows()
        }
    for idx, row in out[out["decision"].astype(str).isin(TCT_EXACT_CODES)].iterrows():
        key = tuple(str(row.get(k) or "") for k in KEYS)
        previous = prior_lookup.get(key)
        current_score = _num(row.get("score"))
        if previous is None:
            state = "NOUVEAU_SIGNAL_TCT"
            previous_decision = None
            previous_score = None
            delta = None
        else:
            previous_decision = str(previous.get("decision") or "")
            previous_score = _num(previous.get("score"))
            delta = current_score - previous_score if current_score is not None and previous_score is not None else None
            if previous_decision != str(row.get("decision") or ""):
                state = "TCT_DECISION_MODIFIEE"
            elif delta is not None and abs(delta) >= 2.0:
                state = "TCT_SCORE_EN_MOUVEMENT"
            else:
                state = "STABLE"
        out.at[idx, "previous_decision"] = previous_decision
        out.at[idx, "previous_score"] = previous_score
        out.at[idx, "score_delta"] = delta
        out.at[idx, "tct_change_state"] = state
    return out


def _groups(source: pd.DataFrame):
    if source.empty:
        return []
    status = source["ci_final_status"].astype(str)
    return [
        ("RECOMMANDATION_TOTALEMENT_VALIDEE", "Recommandations totalement validées", source[status.eq("RECOMMANDATION_TOTALEMENT_VALIDEE")]),
        ("BUY_INTERNE_ATTENTE_SOURCES", "BUY internes en attente de confirmation", source[status.eq("BUY_INTERNE_ATTENTE_SOURCES")]),
        ("TCT_T2_SOURCE_CONFIRMED", "TCT T2 source-confirmé — décision-support SHADOW", source[status.eq("TCT_T2_SOURCE_CONFIRMED")]),
        ("TCT_T2_ATTENTE_SOURCES", "TCT T2 en attente de sources", source[status.eq("TCT_T2_ATTENTE_SOURCES")]),
        ("TCT_T1_SURVEILLANCE", "TCT T1 — surveillance", source[status.eq("TCT_T1_SURVEILLANCE")]),
        ("SURVEILLANCE_INTERNE", "Surveillance / review", source[status.eq("SURVEILLANCE_INTERNE")]),
    ]


def _append_markdown(path: Path, source: pd.DataFrame) -> None:
    lines = [
        "",
        "## 6. Validation finale Boursorama + Investing",
        "",
        "TCT : T1 reste surveillance ; T2 exact source-confirmé reste un signal SHADOW de décision-support et n'est pas requalifié en BUY de production.",
        "",
    ]
    if source.empty:
        lines.append("- Aucune présélection Action/ETF à qualifier.")
    else:
        for _status, title, subset in _groups(source):
            lines.extend([f"### {title}", ""])
            if subset.empty:
                lines.append("- Aucune ligne.")
                continue
            for _, row in subset.iterrows():
                change = str(row.get("tct_change_state") or "")
                change_txt = f" — évolution `{change}`" if change else ""
                lines.append(
                    f"- **{row.get('asset_class')} {row.get('name') or row.get('isin')} [{row.get('horizon')}]** — "
                    f"{row.get('decision')}{change_txt} — gate `{row.get('source_validation_state', 'N/A')}` — "
                    f"Boursorama={'OK' if _bool(row.get('boursorama_priority_ready')) else 'INCOMPLET'} — "
                    f"Investing {row.get('investing_required_timeframe', 'N/A')}={row.get('investing_horizon_signal', 'N/A')} "
                    f"(jour={row.get('investing_daily_signal', 'N/A')}, semaine={row.get('investing_weekly_signal', 'N/A')}, mois={row.get('investing_monthly_signal', 'N/A')})."
                )
                reasons = row.get("source_validation_reasons")
                if pd.notna(reasons) and str(reasons) != "OK":
                    lines.append(f"  - Attente : {reasons}")
    with path.open("a", encoding="utf-8") as handle:
        handle.write("\n".join(lines) + "\n")


def _append_word(path: Path, source: pd.DataFrame) -> None:
    from docx import Document

    document = Document(path)
    document.add_heading("6. Validation finale Boursorama + Investing", level=1)
    document.add_paragraph(
        "TCT : T1 reste surveillance ; T2 exact source-confirmé reste un signal SHADOW de décision-support. "
        "Il n'est pas requalifié en recommandation BUY de production."
    )
    if source.empty:
        document.add_paragraph("Aucune présélection Action/ETF à qualifier.")
    else:
        for _status, title, subset in _groups(source):
            document.add_heading(title, level=2)
            if subset.empty:
                document.add_paragraph("Aucune ligne.")
                continue
            for _, row in subset.iterrows():
                change = str(row.get("tct_change_state") or "")
                document.add_paragraph(
                    f"{row.get('asset_class')} — {row.get('name') or row.get('isin')} — {row.get('horizon')} : "
                    f"décision {row.get('decision')} ; évolution {change or 'n/a'} ; gate {row.get('source_validation_state', 'N/A')} ; "
                    f"Boursorama {'validé' if _bool(row.get('boursorama_priority_ready')) else 'incomplet'} ; "
                    f"Investing requis {row.get('investing_required_timeframe', 'N/A')}=STRONG_BUY, observé {row.get('investing_horizon_signal', 'N/A')} ; "
                    f"Daily {row.get('investing_daily_signal', 'N/A')} / Weekly {row.get('investing_weekly_signal', 'N/A')} / Monthly {row.get('investing_monthly_signal', 'N/A')}.",
                    style="List Bullet",
                )
                reasons = row.get("source_validation_reasons")
                if pd.notna(reasons) and str(reasons) != "OK":
                    document.add_paragraph(f"Attente : {reasons}")
    document.save(path)


def _enrich_matrix(path: Path, source: pd.DataFrame) -> None:
    matrix = _read(path)
    if source.empty:
        return
    keys = KEYS
    keep = keys + [c for c in SOURCE_COLUMNS if c not in keys and c in source] + ["ci_final_status"]
    if matrix.empty:
        matrix = pd.DataFrame(columns=["asset_class", "horizon", "isin", "name", "decision", "score", "action_bucket", "change_state"])
    before = len(matrix)
    matrix = matrix.merge(source[keep].drop_duplicates(keys, keep="last"), on=keys, how="left", sort=False, validate="many_to_one")
    if len(matrix) != before:
        raise RuntimeError("DECISION_BRIEF_SOURCE_JOIN_ROW_COUNT_MUTATION")
    existing_keys = set(zip(matrix.get("asset_class", []), matrix.get("horizon", []), matrix.get("isin", [])))
    additions: list[dict] = []
    for _, row in source[source["decision"].astype(str).isin(TCT_EXACT_CODES)].iterrows():
        key = (row.get("asset_class"), row.get("horizon"), row.get("isin"))
        if key in existing_keys:
            continue
        final_status = str(row.get("ci_final_status") or "")
        bucket = "TCT T2 — ACTION SHADOW" if final_status == "TCT_T2_SOURCE_CONFIRMED" else "TCT — SURVEILLER / ATTENDRE"
        additions.append(
            {
                "asset_class": row.get("asset_class"),
                "horizon": row.get("horizon"),
                "isin": row.get("isin"),
                "name": row.get("name"),
                "decision": row.get("decision"),
                "score": row.get("score"),
                "coverage_pct": row.get("coverage_pct"),
                "action_bucket": bucket,
                "change_state": row.get("tct_change_state") or "TCT_EXACT_CURRENT",
                "previous_decision": row.get("previous_decision"),
                "previous_score": row.get("previous_score"),
                "score_delta": row.get("score_delta"),
                **{column: row.get(column) for column in keep if column not in keys},
            }
        )
    if additions:
        matrix = pd.concat([matrix, pd.DataFrame(additions)], ignore_index=True, sort=False)
    if matrix.duplicated(keys).any():
        raise RuntimeError("DECISION_BRIEF_DUPLICATE_KEYS_AFTER_TCT_APPEND")
    matrix.to_csv(path, sep=";", encoding="utf-8-sig", index=False)


def _enrich_snapshot(root: Path, source: pd.DataFrame) -> None:
    path = root / "state" / "provenance" / "CI_DECISION_SNAPSHOT.csv"
    snapshot = _read(path)
    if source.empty:
        return
    keys = KEYS
    keep = keys + [c for c in ("source_validation_state", "ci_source_eligible", "source_fully_validated", "investing_horizon_signal", "investing_required_timeframe", "boursorama_priority_ready") if c in source]
    if not snapshot.empty:
        before = len(snapshot)
        snapshot = snapshot.merge(source[keep].drop_duplicates(keys, keep="last"), on=keys, how="left", sort=False, validate="many_to_one")
        if len(snapshot) != before:
            raise RuntimeError("DECISION_SNAPSHOT_SOURCE_JOIN_ROW_COUNT_MUTATION")
    else:
        snapshot = pd.DataFrame()
    existing_keys = set(zip(snapshot.get("asset_class", []), snapshot.get("horizon", []), snapshot.get("isin", []))) if not snapshot.empty else set()
    additions = []
    for _, row in source[source["decision"].astype(str).isin(TCT_EXACT_CODES)].iterrows():
        key = (row.get("asset_class"), row.get("horizon"), row.get("isin"))
        if key in existing_keys:
            continue
        additions.append(
            {
                "generated_at_utc": row.get("generated_at_utc"),
                "asset_class": row.get("asset_class"),
                "horizon": row.get("horizon"),
                "isin": row.get("isin"),
                "name": row.get("name"),
                "decision": row.get("decision"),
                "score": row.get("score"),
                "coverage_pct": row.get("coverage_pct"),
                "entry_state": None,
                "position_state": None,
                "data_confidence": "TCT_EXACT_INTERNAL",
                **{column: row.get(column) for column in keep if column not in keys},
            }
        )
    if additions:
        snapshot = pd.concat([snapshot, pd.DataFrame(additions)], ignore_index=True, sort=False)
    if snapshot.duplicated(keys).any():
        raise RuntimeError("DECISION_SNAPSHOT_DUPLICATE_KEYS_AFTER_TCT_APPEND")
    path.parent.mkdir(parents=True, exist_ok=True)
    snapshot.to_csv(path, sep=";", encoding="utf-8-sig", index=False)


def _write_snapshot_frame(path: Path, frame: pd.DataFrame) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    frame.to_csv(path, sep=";", encoding="utf-8-sig", index=False)


def run(root: Path = ROOT) -> dict:
    snapshot_path = root / "state" / "provenance" / "CI_DECISION_SNAPSHOT.csv"
    snapshot_existed = snapshot_path.exists()
    original_snapshot = _read(snapshot_path)
    prior_tct = _tct_prior(original_snapshot)
    if snapshot_existed and not original_snapshot.empty:
        _write_snapshot_frame(snapshot_path, _legacy_snapshot_only(original_snapshot))
    try:
        payload = legacy.run(root)
    except Exception:
        if snapshot_existed:
            _write_snapshot_frame(snapshot_path, original_snapshot)
        elif snapshot_path.exists():
            snapshot_path.unlink()
        raise

    source = _tct_history(_source_rows(root), prior_tct)
    status = source.get("ci_final_status", pd.Series(dtype=str)).astype(str) if not source.empty else pd.Series(dtype=str)
    metrics = {
        "fully_validated_recommendations": int(status.eq("RECOMMANDATION_TOTALEMENT_VALIDEE").sum()),
        "internal_buy_waiting_sources": int(status.eq("BUY_INTERNE_ATTENTE_SOURCES").sum()),
        "tct_t1_surveillance": int(status.eq("TCT_T1_SURVEILLANCE").sum()),
        "tct_t2_waiting_sources": int(status.eq("TCT_T2_ATTENTE_SOURCES").sum()),
        "tct_t2_source_confirmed": int(status.eq("TCT_T2_SOURCE_CONFIRMED").sum()),
        "surveillance_rows": int(status.eq("SURVEILLANCE_INTERNE").sum()),
    }
    payload["version"] = VERSION
    payload.update(metrics)
    payload["source_gate_required_for_entry"] = True
    payload["tct_t2_remains_shadow_decision_support"] = True
    payload["tct_false_removal_guard"] = True
    payload["prior_tct_snapshot_rows"] = int(len(prior_tct))
    payload["source_gate_score_influence"] = 0.0
    payload["source_gate_decision_influence"] = 0.0
    pending_total = metrics["internal_buy_waiting_sources"] + metrics["tct_t2_waiting_sources"]
    if pending_total:
        warnings = list(payload.get("warnings") or [])
        warnings.append(f"{pending_total} dossier(s) d'entrée restent en attente de validation Boursorama/Investing.")
        payload["warnings"] = warnings
        if payload.get("decision_status") == "READY_FOR_REVIEW":
            payload["decision_status"] = "REVIEW_WITH_WARNINGS"

    outdir = root / "outputs" / "decision_brief"
    md_path = outdir / "DECISION_BRIEF.md"
    json_path = outdir / "DECISION_BRIEF.json"
    docx_path = outdir / "CI_DECISION_BRIEF_V3.docx"
    matrix_path = outdir / "CI_DECISION_MATRIX_V3.csv"
    _append_markdown(md_path, source)
    _append_word(docx_path, source)
    _enrich_matrix(matrix_path, source)
    _enrich_snapshot(root, source)
    source.to_csv(outdir / "CI_DECISION_SOURCE_GATE_V21_16.csv", sep=";", encoding="utf-8-sig", index=False)
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    return payload


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", default=str(ROOT))
    args = parser.parse_args()
    print(json.dumps(run(Path(args.root)), ensure_ascii=False, indent=2, default=str))


if __name__ == "__main__":
    main()
