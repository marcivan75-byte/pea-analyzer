from __future__ import annotations

from pathlib import Path
import json

ROOT = Path(__file__).resolve().parents[3]


def _read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}


def _fmt(value, digits: int = 2) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return "N/A"


def _basket_lines(portfolio: dict) -> list[str]:
    if not portfolio or portfolio.get("status") != "OK":
        return [
            "## Panier des recommandations",
            "",
            "- Analyse indisponible : aucun panier actif avec historique suffisant.",
            "- Portefeuille réellement détenu : **NON RENSEIGNÉ**.",
        ]
    return [
        "## Panier des recommandations — distinct du portefeuille détenu",
        "",
        f"- ISIN actifs uniques: {portfolio.get('unique_isins', 'N/A')}",
        f"- Lignes multi-horizons neutralisées: {portfolio.get('duplicate_horizon_rows_removed', 'N/A')}",
        f"- Méthode de pondération diagnostique: {portfolio.get('weight_method', 'N/A')}",
        f"- Bêta 252j du panier: {_fmt(portfolio.get('portfolio_beta_252d'))}",
        f"- Downside bêta 252j du panier: {_fmt(portfolio.get('portfolio_downside_beta_252d'))}",
        f"- Corrélation moyenne: {_fmt(portfolio.get('mean_pair_correlation_252d'))}",
        f"- Corrélation stress: {_fmt(portfolio.get('mean_stress_pair_correlation'))}",
        f"- Concentration/diversification: **{portfolio.get('diversification_warning', 'N/A')}**",
        f"- Moteur économique dominant: {portfolio.get('top_engine', 'N/A')} ({_fmt(portfolio.get('top_engine_share_pct'))}%)",
        "- Portefeuille réellement détenu : **NON RENSEIGNÉ** ; aucun Portfolio Fit réel n'est calculé.",
    ]


def build_markdown(root: Path = ROOT) -> str:
    audit = _read_json(root / "outputs" / "audit" / "BETA_CORRELATION_RISK_ENGINE.json")
    portfolio = _read_json(root / "outputs" / "risk" / "PORTFOLIO_RISK_SUMMARY.json")
    validation = _read_json(root / "config" / "BETA_RISK_ROBUST_VALIDATION_STATUS.json")
    benchmark = audit.get("benchmark", {}) if isinstance(audit.get("benchmark"), dict) else {}
    policy = validation.get("production_policy", {}) if isinstance(validation.get("production_policy"), dict) else {}
    lines = [
        "# RISK V1.1 — CONTROL CENTER CI",
        "",
        f"**Statut moteur:** {audit.get('status', 'NOT_RUN')}  ",
        f"**Benchmark:** {benchmark.get('label', 'N/A')} — {benchmark.get('status', 'N/A')}  ",
        f"**Couverture bêta 252j:** {_fmt(audit.get('coverage_pct'))}%  ",
        "",
    ]
    lines.extend(_basket_lines(portfolio))
    lines.extend(["", "## Stress systématique du panier diagnostique"])
    stress = portfolio.get("systematic_stress_scenarios_pct", {})
    if isinstance(stress, dict) and stress:
        for scenario, estimate in stress.items():
            lines.append(f"- Marché {scenario}% → composante systématique estimée {_fmt(estimate)}%")
    else:
        lines.append("- N/A")
    lines.extend(
        [
            "",
            "## Gouvernance",
            f"- Validation économique: {validation.get('status', 'NOT_RUN')}",
            "- Ce bloc mesure la concentration du panier de recommandations, pas le portefeuille réellement détenu.",
            "- Un même ISIN présent sur plusieurs horizons ne compte qu'une fois dans le panier agrégé.",
            "- Sizing bêta permanent: **REJETÉ**",
            "- Régime V1.1: **REJETÉ**",
            "- Régime V1.2: **REJETÉ**",
            f"- Influence score: {policy.get('score_influence', 0.0)}",
            f"- Influence décision: {policy.get('decision_influence', 0.0)}",
            f"- Influence sizing: {policy.get('sizing_execution_influence', 0.0)}",
            f"- Influence stop: {policy.get('stop_loss_influence', 0.0)}",
            "- Usage autorisé: observabilité, stress, concentration/diversification et warnings uniquement.",
            "- Les scénarios bêta représentent une composante systématique, pas une prévision de perte totale.",
            "",
        ]
    )
    return "\n".join(lines)


def _write_ci_word(path: Path, portfolio: dict, audit: dict) -> None:
    from docx import Document
    from docx.shared import Cm, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading("CI — Risque du panier de recommandations", level=0)
    document.add_paragraph(
        "Ce document complète la restitution CI. Il décrit le risque agrégé des recommandations actives uniques par ISIN. "
        "Il ne représente pas le portefeuille détenu et n'a aucune influence sur les scores, décisions, allocations ou ordres."
    )
    document.add_heading("Synthèse décisionnelle", level=1)
    if portfolio.get("status") != "OK":
        document.add_paragraph("Panier actif insuffisant pour produire une analyse de concentration fiable.")
    else:
        items = [
            f"ISIN actifs uniques : {portfolio.get('unique_isins', 'N/A')}",
            f"Doublons multi-horizons neutralisés : {portfolio.get('duplicate_horizon_rows_removed', 'N/A')}",
            f"Alerte concentration/diversification : {portfolio.get('diversification_warning', 'N/A')}",
            f"Moteur dominant : {portfolio.get('top_engine', 'N/A')} ({_fmt(portfolio.get('top_engine_share_pct'))}%)",
            f"Bêta 252j : {_fmt(portfolio.get('portfolio_beta_252d'))}",
            f"Downside bêta 252j : {_fmt(portfolio.get('portfolio_downside_beta_252d'))}",
            f"Corrélation stress moyenne : {_fmt(portfolio.get('mean_stress_pair_correlation'))}",
        ]
        for item in items:
            document.add_paragraph(item, style="List Bullet")
    document.add_heading("Lecture CI", level=1)
    warning = str(portfolio.get("diversification_warning", "N/A"))
    if warning in {"RED", "ORANGE"}:
        document.add_paragraph(
            "Le panier présente une concentration ou une corrélation de stress élevée. Le CI doit examiner les recommandations en concurrence, "
            "et non additionner mécaniquement plusieurs expositions économiques proches."
        )
    elif warning == "AMBER":
        document.add_paragraph(
            "Le panier présente une concentration intermédiaire. Une vérification des expositions économiques communes est recommandée avant arbitrage."
        )
    else:
        document.add_paragraph(
            "Aucune alerte forte de concentration n'est publiée par le diagnostic actuel. Cette lecture ne remplace pas l'analyse du portefeuille réellement détenu."
        )
    document.add_heading("Portefeuille réellement détenu", level=1)
    document.add_paragraph(
        "NON RENSEIGNÉ. Le Portfolio Fit réel reste volontairement désactivé tant qu'un fichier de positions effectives n'est pas fourni. "
        "Aucun portefeuille équipondéré fictif n'est utilisé à sa place."
    )
    document.add_heading("Contrôles", level=1)
    document.add_paragraph(f"Statut moteur risque : {audit.get('status', 'NOT_RUN')}")
    document.add_paragraph("Influence score/décision/sizing/stop : 0. Aucun ordre réel.")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def run(root: Path = ROOT) -> dict:
    outdir = root / "outputs" / "mobile"
    committee_dir = root / "outputs" / "committee_master"
    outdir.mkdir(parents=True, exist_ok=True)
    committee_dir.mkdir(parents=True, exist_ok=True)
    path = outdir / "RISK_V1_1_CONTROL_CENTER.md"
    text = build_markdown(root)
    path.write_text(text, encoding="utf-8")
    audit = _read_json(root / "outputs" / "audit" / "BETA_CORRELATION_RISK_ENGINE.json")
    portfolio = _read_json(root / "outputs" / "risk" / "PORTFOLIO_RISK_SUMMARY.json")
    word_path = committee_dir / "CI_RISQUE_PANIER_RECOMMANDATIONS.docx"
    _write_ci_word(word_path, portfolio, audit)
    return {
        "status": "SUCCESS",
        "output": str(path.relative_to(root)),
        "ci_word_output": str(word_path.relative_to(root)),
        "analysis_universe": portfolio.get("analysis_universe", "UNAVAILABLE"),
        "is_real_portfolio": False,
        "real_portfolio_fit_status": "NOT_AVAILABLE_NO_PORTFOLIO_INPUT",
        "score_or_decision_mutation": False,
    }


if __name__ == "__main__":
    print(json.dumps(run(), ensure_ascii=False))
