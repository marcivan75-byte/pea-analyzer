from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone
from pathlib import Path
from time import perf_counter
import csv
import io
import json
import math
import os

import pandas as pd
import requests

ROOT = Path(__file__).resolve().parents[3]
VERSION = "DAILY_MARKET_ORIENTATION_V21_15_9"
CACHE_RELATIVE = Path("state/provenance/source_cache/DAILY_MARKET_ORIENTATION_V21_15_9.json")
OUTPUT_RELATIVE = Path("outputs/market_orientation/DAILY_MARKET_ORIENTATION_V21_15_9.json")
CSV_RELATIVE = Path("outputs/market_orientation/DAILY_MARKET_ORIENTATION_V21_15_9.csv")
AUDIT_RELATIVE = Path("outputs/audit/DAILY_MARKET_ORIENTATION_V21_15_9.json")


def _now() -> datetime:
    return datetime.now(timezone.utc)


def _num(value):
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return number if math.isfinite(number) else None


def _pct_change(value, previous):
    value = _num(value)
    previous = _num(previous)
    if value is None or previous in (None, 0.0):
        return None
    return (value / previous - 1.0) * 100.0


def _direction(value, previous, *, inverse: bool = False) -> str:
    value = _num(value)
    previous = _num(previous)
    if value is None or previous is None:
        return "UNKNOWN"
    if math.isclose(value, previous, rel_tol=0.0, abs_tol=1e-12):
        return "NEUTRAL"
    rising = value > previous
    if inverse:
        return "RISK_OFF" if rising else "RISK_ON"
    return "RISK_ON" if rising else "RISK_OFF"


def _load_cache(root: Path) -> dict:
    path = root / CACHE_RELATIVE
    if not path.exists():
        return {"version": VERSION, "indicators": {}}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError, json.JSONDecodeError):
        return {"version": VERSION, "indicators": {}}
    if payload.get("version") != VERSION or not isinstance(payload.get("indicators"), dict):
        return {"version": VERSION, "indicators": {}}
    return payload


def _save_cache(root: Path, payload: dict) -> None:
    path = root / CACHE_RELATIVE
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    tmp.replace(path)


def _get_json(url: str, *, params: dict | None = None, headers: dict | None = None, timeout: float = 5.0):
    response = requests.get(url, params=params, headers=headers, timeout=timeout)
    response.raise_for_status()
    return response.json()


def _fetch_vixcls() -> dict:
    api_key = str(os.environ.get("FRED_API_KEY") or "").strip()
    source_url = "https://fred.stlouisfed.org/series/VIXCLS"
    observations: list[tuple[str, float]] = []
    if api_key:
        payload = _get_json(
            "https://api.stlouisfed.org/fred/series/observations",
            params={
                "series_id": "VIXCLS",
                "api_key": api_key,
                "file_type": "json",
                "sort_order": "desc",
                "limit": 10,
            },
        )
        for row in payload.get("observations", []):
            value = _num(row.get("value"))
            if value is not None:
                observations.append((str(row.get("date") or ""), value))
    else:
        response = requests.get("https://fred.stlouisfed.org/graph/fredgraph.csv?id=VIXCLS", timeout=5.0)
        response.raise_for_status()
        parsed = list(csv.DictReader(io.StringIO(response.text)))
        for row in reversed(parsed[-20:]):
            value = _num(row.get("VIXCLS"))
            if value is not None:
                observations.append((str(row.get("DATE") or row.get("observation_date") or ""), value))
    if not observations:
        raise RuntimeError("VIXCLS_NO_OBSERVATION")
    latest_date, latest = observations[0]
    previous = observations[1][1] if len(observations) > 1 else None
    return {
        "indicator": "VIXCLS",
        "market": "US",
        "value": latest,
        "previous": previous,
        "change_abs": latest - previous if previous is not None else None,
        "change_pct": _pct_change(latest, previous),
        "rating": "",
        "risk_direction": _direction(latest, previous, inverse=True),
        "as_of": latest_date,
        "source": "FRED / CBOE",
        "source_url": source_url,
        "status": "LIVE",
    }


def _fetch_cnn_fear_greed() -> dict:
    source_url = "https://www.cnn.com/markets/fear-and-greed"
    payload = _get_json(
        "https://production.dataviz.cnn.io/index/fearandgreed/graphdata",
        headers={
            "User-Agent": "Mozilla/5.0 (compatible; PEA-Analyzer/21.15; daily-market-context)",
            "Accept": "application/json",
            "Origin": "https://edition.cnn.com",
            "Referer": "https://edition.cnn.com/",
        },
    )
    block = payload.get("fear_and_greed") or {}
    score = _num(block.get("score"))
    previous = _num(block.get("previous_close"))
    if score is None:
        raise RuntimeError("CNN_FEAR_GREED_SCORE_MISSING")
    return {
        "indicator": "CNN_FEAR_GREED",
        "market": "US_SENTIMENT",
        "value": score,
        "previous": previous,
        "change_abs": score - previous if previous is not None else None,
        "change_pct": _pct_change(score, previous),
        "rating": str(block.get("rating") or "").strip().upper(),
        "risk_direction": _direction(score, previous, inverse=False),
        "as_of": str(block.get("timestamp") or ""),
        "source": "CNN Fear & Greed",
        "source_url": source_url,
        "status": "LIVE",
        "previous_1_week": _num(block.get("previous_1_week")),
        "previous_1_month": _num(block.get("previous_1_month")),
    }


def _parse_stoxx_date(value: str) -> datetime | None:
    text = str(value or "").strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(text, fmt).replace(tzinfo=timezone.utc)
        except ValueError:
            continue
    return None


def _fetch_vstoxx() -> dict:
    data_url = "https://www.stoxx.com/document/Indices/Current/HistoricalData/h_v2tx.txt"
    response = requests.get(
        data_url,
        headers={"User-Agent": "Mozilla/5.0 (compatible; PEA-Analyzer/21.15; daily-market-context)"},
        timeout=5.0,
    )
    response.raise_for_status()
    valid: list[tuple[datetime, float]] = []
    for raw in response.text.splitlines():
        line = raw.strip()
        if not line or not line[0].isdigit():
            continue
        delimiter = ";" if ";" in line else ","
        parts = [part.strip() for part in line.split(delimiter)]
        if len(parts) < 2:
            continue
        observed = _parse_stoxx_date(parts[0])
        if observed is None:
            continue
        value = None
        if delimiter == ";" and len(parts) >= 3 and parts[1].upper() == "V2TX":
            value = _num(parts[2])
        elif delimiter == ",":
            value = _num(parts[1])
        if value is not None:
            valid.append((observed, value))
    if not valid:
        raise RuntimeError("VSTOXX_OFFICIAL_HISTORY_MISSING")
    valid.sort(key=lambda item: item[0])
    observed, value = valid[-1]
    previous = valid[-2][1] if len(valid) > 1 else None
    age_days = (_now().date() - observed.date()).days
    if age_days < 0 or age_days > 10:
        raise RuntimeError(f"VSTOXX_OFFICIAL_HISTORY_STALE:{observed.date().isoformat()}:age_days={age_days}")
    return {
        "indicator": "VSTOXX",
        "market": "EUROPE",
        "value": value,
        "previous": previous,
        "change_abs": value - previous if previous is not None else None,
        "change_pct": _pct_change(value, previous),
        "rating": "",
        "risk_direction": _direction(value, previous, inverse=True),
        "as_of": observed.date().isoformat(),
        "source": "STOXX official V2TX historical data",
        "source_url": data_url,
        "status": "LIVE",
        "symbol": "V2TX",
        "freshness_age_days": age_days,
    }


def _fallback(name: str, cached: dict | None, exc: Exception) -> dict:
    if cached:
        row = dict(cached)
        row["status"] = "CACHE_FALLBACK"
        row["live_error_type"] = type(exc).__name__
        row["live_error"] = str(exc)[:240]
        return row
    return {
        "indicator": name,
        "value": None,
        "previous": None,
        "change_abs": None,
        "change_pct": None,
        "rating": "",
        "risk_direction": "UNKNOWN",
        "as_of": "",
        "source": "",
        "source_url": "",
        "status": "UNAVAILABLE",
        "live_error_type": type(exc).__name__,
        "live_error": str(exc)[:240],
    }


def _orientation(indicators: list[dict]) -> tuple[str, str]:
    directions = [str(row.get("risk_direction") or "UNKNOWN") for row in indicators]
    risk_on = directions.count("RISK_ON")
    risk_off = directions.count("RISK_OFF")
    if risk_on >= 2 and risk_on > risk_off:
        orientation = "SUPPORTIVE"
    elif risk_off >= 2 and risk_off > risk_on:
        orientation = "CAUTION"
    else:
        orientation = "MIXED_NEUTRAL"
    detail = ", ".join(
        f"{row.get('indicator')}={row.get('risk_direction')}"
        + (f" ({row.get('rating')})" if row.get("rating") else "")
        for row in indicators
    )
    return orientation, detail


def _write_outputs(root: Path, payload: dict) -> None:
    json_path = root / OUTPUT_RELATIVE
    csv_path = root / CSV_RELATIVE
    audit_path = root / AUDIT_RELATIVE
    json_path.parent.mkdir(parents=True, exist_ok=True)
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    audit_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    pd.DataFrame(payload.get("indicators") or []).to_csv(csv_path, sep=";", index=False, encoding="utf-8-sig")


def run(root: Path = ROOT) -> dict:
    started = perf_counter()
    cache = _load_cache(root)
    cached = cache.get("indicators") or {}
    tasks = {
        "VIXCLS": _fetch_vixcls,
        "CNN_FEAR_GREED": _fetch_cnn_fear_greed,
        "VSTOXX": _fetch_vstoxx,
    }
    results: dict[str, dict] = {}
    with ThreadPoolExecutor(max_workers=3, thread_name_prefix="daily-market-orientation") as pool:
        future_map = {pool.submit(fn): name for name, fn in tasks.items()}
        for future in as_completed(future_map):
            name = future_map[future]
            try:
                results[name] = future.result()
            except Exception as exc:
                results[name] = _fallback(name, cached.get(name), exc)

    ordered = [results[name] for name in ("VIXCLS", "CNN_FEAR_GREED", "VSTOXX")]
    orientation, orientation_detail = _orientation(ordered)
    live_count = sum(str(row.get("status")) == "LIVE" for row in ordered)
    usable_count = sum(row.get("value") is not None for row in ordered)
    payload = {
        "status": "SUCCESS" if usable_count == 3 else ("SUCCESS_WITH_CACHE_OR_PARTIAL" if usable_count else "UNAVAILABLE_NON_BLOCKING"),
        "version": VERSION,
        "generated_at_utc": _now().isoformat(),
        "scope": "DAILY_UPSTREAM_LIGHT_MARKET_ORIENTATION",
        "orientation": orientation,
        "orientation_detail": orientation_detail,
        "indicators": ordered,
        "live_indicators": int(live_count),
        "usable_indicators": int(usable_count),
        "decision_influence": False,
        "score_influence": 0.0,
        "weights_changed": False,
        "thresholds_changed": False,
        "criteria_changed": False,
        "can_create_buy": False,
        "can_block_buy": False,
        "real_orders_enabled": False,
        "elapsed_seconds": round(perf_counter() - started, 6),
    }
    cache_payload = {
        "version": VERSION,
        "updated_at_utc": payload["generated_at_utc"],
        "indicators": {
            row["indicator"]: row
            for row in ordered
            if row.get("value") is not None and str(row.get("status")) == "LIVE"
        },
    }
    if not cache_payload["indicators"]:
        cache_payload["indicators"] = cached
    else:
        for key, value in cached.items():
            cache_payload["indicators"].setdefault(key, value)
    _save_cache(root, cache_payload)
    _write_outputs(root, payload)
    return payload


def _fmt(value, digits: int = 2) -> str:
    number = _num(value)
    return "n/a" if number is None else f"{number:.{digits}f}"


def _market_lines(payload: dict) -> list[str]:
    lines = [
        "## Orientation marché légère — contexte amont",
        "",
        f"Orientation synthétique : **{payload.get('orientation', 'n/a')}** — contexte uniquement, influence score/décision = 0.",
        "",
        "| Indicateur | Valeur | Variation | Lecture | As of | Statut |",
        "|---|---:|---:|---|---|---|",
    ]
    for row in payload.get("indicators") or []:
        change = _fmt(row.get("change_pct"))
        if change != "n/a":
            change += "%"
        reading = str(row.get("rating") or row.get("risk_direction") or "n/a")
        lines.append(
            f"| {row.get('indicator')} | {_fmt(row.get('value'))} | {change} | {reading} | {row.get('as_of') or 'n/a'} | {row.get('status') or 'n/a'} |"
        )
    lines.extend(["", f"Détail : {payload.get('orientation_detail', '')}", ""])
    return lines


def publish_ci_context(root: Path, payload: dict) -> dict:
    committee = root / "outputs" / "committee_master"
    mobile = root / "outputs" / "mobile"
    committee.mkdir(parents=True, exist_ok=True)
    mobile.mkdir(parents=True, exist_ok=True)
    indicators = pd.DataFrame(payload.get("indicators") or [])
    indicators.to_csv(committee / "CI_MARKET_ORIENTATION.csv", sep=";", index=False, encoding="utf-8-sig")

    results = {"csv": "outputs/committee_master/CI_MARKET_ORIENTATION.csv"}
    android_path = mobile / "ANDROID_CI_CONTROL_CENTER.md"
    if android_path.exists():
        text = android_path.read_text(encoding="utf-8")
        android_path.write_text("\n".join(_market_lines(payload)) + "\n" + text, encoding="utf-8")
        results["android"] = str(android_path.relative_to(root))

    word_path = committee / "CI_COMITE_INVESTISSEMENT.docx"
    if word_path.exists():
        try:
            from docx import Document
            document = Document(word_path)
            document.add_heading("Orientation marché légère — contexte amont", level=1)
            document.add_paragraph(
                f"Orientation synthétique : {payload.get('orientation', 'n/a')}. "
                "Bloc contextuel uniquement : influence sur score et décision = 0."
            )
            table = document.add_table(rows=1, cols=6)
            headers = ["Indicateur", "Valeur", "Variation %", "Lecture", "As of", "Statut"]
            for idx, label in enumerate(headers):
                table.rows[0].cells[idx].text = label
            for row in payload.get("indicators") or []:
                cells = table.add_row().cells
                cells[0].text = str(row.get("indicator") or "")
                cells[1].text = _fmt(row.get("value"))
                cells[2].text = _fmt(row.get("change_pct"))
                cells[3].text = str(row.get("rating") or row.get("risk_direction") or "")
                cells[4].text = str(row.get("as_of") or "")
                cells[5].text = str(row.get("status") or "")
            document.save(word_path)
            results["word"] = str(word_path.relative_to(root))
        except Exception as exc:
            results["word_error"] = f"{type(exc).__name__}:{str(exc)[:160]}"

    excel_path = committee / "CI_REFERENTIEL_PONDERE.xlsx"
    if excel_path.exists():
        try:
            from openpyxl import load_workbook
            workbook = load_workbook(excel_path)
            if "MARKET_ORIENTATION" in workbook.sheetnames:
                del workbook["MARKET_ORIENTATION"]
            sheet = workbook.create_sheet("MARKET_ORIENTATION", 0)
            sheet.append(["Orientation", payload.get("orientation"), "Influence décision", 0, "Influence score", 0])
            sheet.append([])
            columns = ["indicator", "market", "value", "previous", "change_abs", "change_pct", "rating", "risk_direction", "as_of", "source", "source_url", "status"]
            sheet.append(columns)
            for row in payload.get("indicators") or []:
                sheet.append([row.get(column) for column in columns])
            workbook.save(excel_path)
            results["excel"] = str(excel_path.relative_to(root))
        except Exception as exc:
            results["excel_error"] = f"{type(exc).__name__}:{str(exc)[:160]}"
    return results
