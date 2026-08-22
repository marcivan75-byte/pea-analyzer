from __future__ import annotations

from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
import argparse
import csv
import json

import pandas as pd


ROOT = Path(__file__).resolve().parents[3]
VERSION = "CI_DECISION_BRIEF_V3"
SELECTED_DECISIONS = {"BUY_CANDIDATE", "WATCH", "WATCH_NOT_TOP2", "REVIEW"}
KEYS = ["asset_class", "horizon", "isin"]
SNAPSHOT_COLUMNS = [
    "generated_at_utc",
    "asset_class",
    "horizon",
    "isin",
    "name",
    "decision",
    "score",
    "coverage_pct",
    "consensus_score",
    "consensus_delta_4w",
    "target_upside_pct",
    "entry_state",
    "position_state",
    "catalyst_state",
    "data_confidence",
]


def _json(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def _rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    try:
        with path.open(encoding="utf-8-sig", newline="") as handle:
            return list(csv.DictReader(handle, delimiter=";"))
    except OSError:
        return []


def _frame(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_csv(path, sep=";", encoding="utf-8-sig", low_memory=False)
    except (OSError, ValueError):
        return pd.DataFrame()


def _reference(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_excel(path, sheet_name="Referentiel_pondere")
    except (OSError, ValueError, ImportError):
        return pd.DataFrame()


def _number(value) -> float | None:
    try:
        result = float(value)
    except (TypeError, ValueError):
        return None
    return result if pd.notna(result) else None


def _text(value, default: str = "n/a") -> str:
    if value is None:
        return default
    try:
        missing = bool(pd.isna(value))
    except (TypeError, ValueError):
        missing = False
    if missing:
        return default
    result = str(value).strip()
    return result if result and result.lower() not in {"nan", "none"} else default


def _safe_bool(value) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"1", "true", "yes", "oui"}


def _step_health(unified: dict) -> tuple[list[str], list[str]]:
    steps = unified.get("steps") if isinstance(unified.get("steps"), dict) else {}
    failed = sorted(name for name, value in steps.items() if isinstance(value, dict) and value.get("status") == "FAILED")
    skipped = sorted(
        name
        for name, value in steps.items()
        if isinstance(value, dict) and str(value.get("status", "")).startswith("SKIPPED_DEPENDENCY")
    )
    return failed, skipped


def _criterion_value(reference: pd.DataFrame, isin: str, horizon: str, names: tuple[str, ...]) -> float | None:
    if reference.empty:
        return None
    required = {"isin", "horizon", "criterion", "raw_value"}
    if not required.issubset(reference.columns):
        return None
    subset = reference[
        (reference["isin"].astype(str) == str(isin))
        & (reference["horizon"].astype(str) == str(horizon))
        & (reference["criterion"].astype(str).isin(names))
    ]
    if subset.empty:
        return None
    for value in subset["raw_value"]:
        parsed = _number(value)
        if parsed is not None:
            return parsed
    return None


def _factor_summary(reference: pd.DataFrame, isin: str, horizon: str, positive: bool) -> str:
    required = {"isin", "horizon", "criterion", "criterion_status", "criterion_score_0_100", "weighted_contribution_points"}
    if reference.empty or not required.issubset(reference.columns):
        return "n/a"
    subset = reference[
        (reference["isin"].astype(str) == str(isin))
        & (reference["horizon"].astype(str) == str(horizon))
        & (reference["criterion_status"].astype(str) == "ACTIVE")
    ].copy()
    if subset.empty:
        return "n/a"
    subset["criterion_score_0_100"] = pd.to_numeric(subset["criterion_score_0_100"], errors="coerce")
    subset["weighted_contribution_points"] = pd.to_numeric(subset["weighted_contribution_points"], errors="coerce")
    subset = subset.dropna(subset=["criterion_score_0_100"])
    if subset.empty:
        return "n/a"
    if positive:
        subset = subset.sort_values(["weighted_contribution_points", "criterion_score_0_100"], ascending=[False, False])
    else:
        subset = subset.sort_values(["criterion_score_0_100", "weighted_contribution_points"], ascending=[True, True])
    parts = []
    for _, row in subset.head(3).iterrows():
        score = _number(row.get("criterion_score_0_100"))
        label = _text(row.get("criterion"))
        parts.append(f"{label}={score:.0f}" if score is not None else label)
    return ", ".join(parts) if parts else "n/a"


def _evidence_quality(reference: pd.DataFrame, isin: str, horizon: str, coverage: float | None) -> tuple[str, str]:
    required = {"isin", "horizon", "criterion_status", "effective_weight_pct"}
    if reference.empty or not required.issubset(reference.columns):
        return "NON_DOCUMENTEE", "référentiel pondéré indisponible"
    subset = reference[
        (reference["isin"].astype(str) == str(isin))
        & (reference["horizon"].astype(str) == str(horizon))
    ].copy()
    if subset.empty:
        return "NON_DOCUMENTEE", "aucune ligne de référentiel pour cette sélection"

    subset["effective_weight_pct"] = pd.to_numeric(subset["effective_weight_pct"], errors="coerce").fillna(0.0)
    active = subset[subset["criterion_status"].astype(str) == "ACTIVE"].copy()
    missing = int((subset["criterion_status"].astype(str) == "MISSING").sum())
    if active.empty:
        return "FAIBLE", f"aucun critère actif documenté ; {missing} manquant(s)"

    source = active["source"] if "source" in active.columns else pd.Series(index=active.index, dtype=object)
    as_of = active["as_of"] if "as_of" in active.columns else pd.Series(index=active.index, dtype=object)
    source_ok = ~source.astype(str).str.lower().isin({"", "nan", "none", "n/a"})
    asof_ok = ~as_of.astype(str).str.lower().isin({"", "nan", "none", "n/a"})
    documented_weight = float(active.loc[source_ok & asof_ok, "effective_weight_pct"].sum())
    validation = active["validation_status"] if "validation_status" in active.columns else pd.Series(index=active.index, dtype=object)
    validation = validation.astype(str).str.upper()
    bad_validation = int(validation.str.contains("FAIL|INVALID|REJECT|STALE", regex=True, na=False).sum())
    cov = coverage if coverage is not None else 0.0

    if cov >= 90.0 and documented_weight >= 70.0 and bad_validation == 0:
        label = "FORTE"
    elif cov >= 75.0 and documented_weight >= 40.0:
        label = "MOYENNE"
    else:
        label = "FAIBLE"
    reason = (
        f"couverture {cov:.1f}% ; poids avec source+date {documented_weight:.1f}% ; "
        f"{missing} critère(s) manquant(s) ; {bad_validation} validation(s) dégradée(s)"
    )
    return label, reason


def _context_lookup(frame: pd.DataFrame, keys: list[str]) -> dict[tuple[str, ...], dict]:
    if frame.empty or any(key not in frame.columns for key in keys):
        return {}
    frame = frame.drop_duplicates(keys, keep="last")
    return {tuple(str(row.get(key, "")) for key in keys): row.to_dict() for _, row in frame.iterrows()}


def _catalyst_lookup(frame: pd.DataFrame) -> dict[str, dict]:
    if frame.empty or "isin" not in frame.columns:
        return {}
    ranked = frame.copy()
    if "movement_potential_score" in ranked.columns:
        ranked["_move"] = pd.to_numeric(ranked["movement_potential_score"], errors="coerce")
        ranked = ranked.sort_values("_move", ascending=False, na_position="last")
    ranked = ranked.drop_duplicates("isin", keep="first")
    return {str(row.get("isin", "")): row.to_dict() for _, row in ranked.iterrows()}


def _prior_lookup(snapshot: pd.DataFrame) -> dict[tuple[str, str, str], dict]:
    return _context_lookup(snapshot, KEYS)


def _delta(current: float | None, previous: float | None) -> float | None:
    if current is None or previous is None:
        return None
    return current - previous


def _change_state(previous: dict | None, decision: str, score_delta: float | None, coverage_delta: float | None) -> str:
    if previous is None:
        return "NOUVELLE_SELECTION"
    if _text(previous.get("decision"), "") != decision:
        return "DECISION_MODIFIEE"
    if score_delta is not None and abs(score_delta) >= 2.0:
        return "SCORE_EN_MOUVEMENT"
    if coverage_delta is not None and abs(coverage_delta) >= 5.0:
        return "COUVERTURE_EN_MOUVEMENT"
    return "STABLE"


def _action_bucket(decision: str, entry_state: str, position_state: str) -> str:
    pos = position_state.upper()
    entry = entry_state.upper()
    if "EXIT" in pos or "EMERGENCY" in pos:
        return "SORTIR / EVITER"
    if "PROTECT" in pos:
        return "CONSERVER / PROTEGER"
    if decision == "BUY_CANDIDATE":
        if any(token in entry for token in ("READY", "CONFIRM", "ELIGIBLE", "TRIGGER", "OPEN", "ENTER")):
            return "ACTION IMMEDIATE"
        return "ATTENDRE DECLENCHEUR"
    if decision in {"WATCH", "WATCH_NOT_TOP2"}:
        return "SURVEILLER"
    if decision == "REVIEW":
        return "REEXAMINER / EVITER"
    return "REVUE CI"


def _why_now(row: dict) -> str:
    parts: list[str] = []
    if row["change_state"] == "NOUVELLE_SELECTION":
        parts.append("nouvelle entrée dans la sélection")
    elif row["change_state"] == "DECISION_MODIFIEE":
        parts.append(f"décision passée de {row['previous_decision']} à {row['decision']}")
    score_delta = row.get("score_delta")
    if score_delta is not None:
        parts.append(f"score S/S-1 {score_delta:+.1f} pt")
    consensus_delta = row.get("consensus_delta_4w")
    if consensus_delta is not None:
        parts.append(f"révision consensus 4 semaines {consensus_delta:+.1f}")
    target = row.get("target_upside_pct")
    if target is not None:
        parts.append(f"potentiel cible {target:+.1f}%")
    catalyst = _text(row.get("catalyst_state"), "")
    if catalyst:
        parts.append(f"contexte dernière minute {catalyst} (shadow)")
    if not parts:
        parts.append("aucun changement majeur détecté ; décision maintenue par le moteur")
    return " ; ".join(parts)


def _invalidation(row: dict) -> str:
    parts: list[str] = []
    for field in ("risk_verdict", "valuation_warning", "correction_alert"):
        value = _text(row.get(field), "")
        if value:
            parts.append(value)
    entry_reasons = _text(row.get("entry_reasons"), "")
    position_reasons = _text(row.get("position_reasons"), "")
    if entry_reasons:
        parts.append(f"entrée: {entry_reasons}")
    if position_reasons:
        parts.append(f"position: {position_reasons}")
    if _safe_bool(row.get("news_technical_conflict")):
        parts.append("conflit news/technique dernière minute")
    exit_state = _text(row.get("exit_state"), "")
    if exit_state:
        parts.append(f"catalyseur: {exit_state}")
    return " | ".join(dict.fromkeys(parts)) if parts else "aucune invalidation contextuelle publiée"


def _decision_rows(decisions: list[dict[str, str]], reference: pd.DataFrame, root: Path) -> list[dict]:
    entry = _frame(root / "outputs" / "committee_master" / "V21_8_ENTRY_EXIT_CHALLENGER.csv")
    sector = _frame(root / "outputs" / "committee_master" / "COMMITTEE_SECTOR_ROTATION_V2_CONTEXT.csv")
    risk = _frame(root / "outputs" / "risk" / "BETA_CORRELATION_RISK_ROWS.csv")
    catalyst = _frame(root / "outputs" / "daily_tct_ct" / "TCT_NEXT_SESSION_CATALYST_V24_4_2.csv")
    previous = _frame(root / "state" / "provenance" / "CI_DECISION_SNAPSHOT.csv")

    entry_lookup = _context_lookup(entry, KEYS)
    sector_lookup = _context_lookup(sector, KEYS)
    risk_lookup = _context_lookup(risk, ["isin"])
    catalyst_lookup = _catalyst_lookup(catalyst)
    prior_lookup = _prior_lookup(previous)

    selected = [row for row in decisions if str(row.get("decision", "")) in SELECTED_DECISIONS]
    output: list[dict] = []
    for raw in selected:
        asset = str(raw.get("asset_class", ""))
        horizon = str(raw.get("horizon", ""))
        isin = str(raw.get("isin", ""))
        key = (asset, horizon, isin)
        ctx = entry_lookup.get(key, {})
        sec = sector_lookup.get(key, {})
        rsk = risk_lookup.get((isin,), {})
        cat = catalyst_lookup.get(isin, {})
        prior = prior_lookup.get(key)
        score = _number(raw.get("score"))
        coverage = _number(raw.get("coverage_pct"))
        previous_score = _number(prior.get("score")) if prior else None
        previous_coverage = _number(prior.get("coverage_pct")) if prior else None
        score_delta = _delta(score, previous_score)
        coverage_delta = _delta(coverage, previous_coverage)
        decision = str(raw.get("decision", ""))
        confidence, confidence_reason = _evidence_quality(reference, isin, horizon, coverage)
        entry_state = _text(ctx.get("v21_8_entry_state"))
        position_state = _text(ctx.get("v21_8_position_state"))
        row = {
            "asset_class": asset,
            "horizon": horizon,
            "isin": isin,
            "name": raw.get("name") or isin,
            "decision": decision,
            "score": score,
            "coverage_pct": coverage,
            "previous_decision": _text(prior.get("decision"), "n/a") if prior else "n/a",
            "previous_score": previous_score,
            "score_delta": score_delta,
            "previous_coverage_pct": previous_coverage,
            "coverage_delta": coverage_delta,
            "change_state": _change_state(prior, decision, score_delta, coverage_delta),
            "consensus_score": _criterion_value(reference, isin, horizon, ("consensus_score_100_v21", "consensus_score_100")),
            "consensus_delta_4w": _criterion_value(reference, isin, horizon, ("consensus_delta_4w",)),
            "target_upside_pct": _criterion_value(reference, isin, horizon, ("target_upside_pct_v21", "target_upside_pct")),
            "positive_factors": _factor_summary(reference, isin, horizon, True),
            "weak_factors": _factor_summary(reference, isin, horizon, False),
            "data_confidence": confidence,
            "data_confidence_reason": confidence_reason,
            "entry_state": entry_state,
            "position_state": position_state,
            "entry_reasons": ctx.get("v21_8_entry_reasons"),
            "position_reasons": ctx.get("v21_8_position_reasons"),
            "valuation_warning": sec.get("valuation_warning"),
            "correction_alert": sec.get("correction_alert"),
            "risk_verdict": rsk.get("risk_verdict"),
            "risk_metric_status": rsk.get("risk_metric_status"),
            "risk_beta_reliability": rsk.get("risk_beta_reliability"),
            "catalyst_state": cat.get("catalyst_state"),
            "movement_potential_score": _number(cat.get("movement_potential_score")),
            "direction_bias_score": _number(cat.get("direction_bias_score")),
            "news_technical_conflict": cat.get("news_technical_conflict"),
            "exit_state": cat.get("exit_state"),
            "data_quality_state": cat.get("data_quality_state"),
            "news_event_types": cat.get("news_event_types"),
        }
        row["action_bucket"] = _action_bucket(decision, entry_state, position_state)
        row["why_now"] = _why_now(row)
        row["invalidation"] = _invalidation(row)
        output.append(row)
    return output


def _removed_rows(root: Path, current_rows: list[dict]) -> list[dict]:
    previous = _frame(root / "state" / "provenance" / "CI_DECISION_SNAPSHOT.csv")
    if previous.empty:
        return []
    current_keys = {(str(row["asset_class"]), str(row["horizon"]), str(row["isin"])) for row in current_rows}
    removed: list[dict] = []
    for _, row in previous.iterrows():
        key = (str(row.get("asset_class", "")), str(row.get("horizon", "")), str(row.get("isin", "")))
        if key not in current_keys:
            removed.append(
                {
                    "asset_class": key[0],
                    "horizon": key[1],
                    "isin": key[2],
                    "name": row.get("name") or key[2],
                    "previous_decision": row.get("decision"),
                    "previous_score": _number(row.get("score")),
                    "change_state": "SORTIE_SELECTION",
                }
            )
    return removed


def _comparisons(rows: list[dict]) -> list[dict]:
    if not rows:
        return []
    frame = pd.DataFrame(rows)
    frame["_score"] = pd.to_numeric(frame["score"], errors="coerce")
    result: list[dict] = []
    for (asset, horizon), group in frame.groupby(["asset_class", "horizon"], sort=False):
        group = group.sort_values("_score", ascending=False, na_position="last").head(3)
        result.append(
            {
                "asset_class": asset,
                "horizon": horizon,
                "candidates": [
                    {
                        "name": row.get("name"),
                        "isin": row.get("isin"),
                        "decision": row.get("decision"),
                        "score": _number(row.get("score")),
                        "coverage_pct": _number(row.get("coverage_pct")),
                        "data_confidence": row.get("data_confidence"),
                        "target_upside_pct": _number(row.get("target_upside_pct")),
                        "consensus_delta_4w": _number(row.get("consensus_delta_4w")),
                        "positive_factors": row.get("positive_factors"),
                        "weak_factors": row.get("weak_factors"),
                        "action_bucket": row.get("action_bucket"),
                    }
                    for _, row in group.iterrows()
                ],
            }
        )
    return result


def _fmt(value, suffix: str = "") -> str:
    parsed = _number(value)
    return "n/a" if parsed is None else f"{parsed:.1f}{suffix}"


def _fmt_delta(value, suffix: str = "") -> str:
    parsed = _number(value)
    return "n/a" if parsed is None else f"{parsed:+.1f}{suffix}"


def _markdown(payload: dict) -> str:
    status_icon = {"READY_FOR_REVIEW": "🟢", "REVIEW_WITH_WARNINGS": "🟠", "BLOCKED": "🔴"}[payload["decision_status"]]
    lines = [
        "# Comité d’investissement — Brief décisionnel V3",
        "",
        f"## {status_icon} Statut : {payload['decision_status']}",
        "",
        f"- Run : `{payload['run_id']}`",
        f"- Sélections à examiner : **{payload['selected_count']}**",
        f"- Changements depuis S-1 : **{payload['changed_count']}** ; sorties de sélection : **{len(payload['removed_from_selection'])}**",
        "- Qualité de preuve = couverture/provenance/fraîcheur documentaire ; **ce n’est pas un nouveau score financier**.",
        "- PREOPEN/POSTMARKET V24.4.2 = contexte SHADOW de dernière minute ; influence score/décision = 0.",
        "- Ordres réels : **DÉSACTIVÉS**.",
        "",
        "## 1. Décisions à prendre",
        "",
        "| Priorité CI | Actif | Horizon | Instrument | Décision | Score | Δ S/S-1 | Preuve |",
        "|---|---|---|---|---|---:|---:|---|",
    ]
    priority = {"ACTION IMMEDIATE": 0, "ATTENDRE DECLENCHEUR": 1, "CONSERVER / PROTEGER": 2, "SURVEILLER": 3, "REEXAMINER / EVITER": 4, "SORTIR / EVITER": 5, "REVUE CI": 6}
    ordered = sorted(payload["decision_rows"], key=lambda row: (priority.get(str(row.get("action_bucket")), 9), -(_number(row.get("score")) or -1.0)))
    if ordered:
        for row in ordered:
            label = str(row.get("name") or row.get("isin") or "n/a").replace("|", "/")
            lines.append(
                f"| {row['action_bucket']} | {row['asset_class']} | {row['horizon']} | {label} | {row['decision']} | "
                f"{_fmt(row.get('score'))} | {_fmt_delta(row.get('score_delta'))} | {row['data_confidence']} |"
            )
    else:
        lines.append("| — | — | — | Aucune sélection publiée | — | — | — | — |")

    lines.extend(["", "## 2. Changements depuis S-1", ""])
    changes = [row for row in ordered if row.get("change_state") != "STABLE"]
    if changes:
        for row in changes:
            lines.append(
                f"- **{row['name']} — {row['asset_class']} {row['horizon']}** : {row['change_state']} ; "
                f"décision {row['previous_decision']} → {row['decision']} ; score {_fmt(row.get('previous_score'))} → {_fmt(row.get('score'))} "
                f"({_fmt_delta(row.get('score_delta'))} pt)."
            )
    else:
        lines.append("- Aucun changement matériel parmi les sélections courantes.")
    for row in payload["removed_from_selection"]:
        lines.append(
            f"- **{row['name']} — {row['asset_class']} {row['horizon']}** : SORTIE_SELECTION ; "
            f"ancienne décision {row.get('previous_decision')}, ancien score {_fmt(row.get('previous_score'))}."
        )

    lines.extend(["", "## 3. Pourquoi maintenant / invalidation", ""])
    focus = [row for row in ordered if row.get("action_bucket") in {"ACTION IMMEDIATE", "ATTENDRE DECLENCHEUR", "CONSERVER / PROTEGER", "SORTIR / EVITER"} or row.get("change_state") != "STABLE"]
    if not focus:
        lines.append("- Aucun dossier nécessitant un commentaire d’exception ; les WATCH stables restent synthétisés dans le tableau principal.")
    for row in focus:
        lines.extend(
            [
                f"### {row['name']} — {row['asset_class']} {row['horizon']}",
                f"- **Pourquoi maintenant :** {row['why_now']}",
                f"- **Facteurs moteurs :** {row['positive_factors']}",
                f"- **Point(s) faible(s) :** {row['weak_factors']}",
                f"- **Invalidation / vigilance :** {row['invalidation']}",
                f"- **Qualité de preuve :** {row['data_confidence']} — {row['data_confidence_reason']}",
            ]
        )

    lines.extend(["", "## 4. Comparatif des meilleurs dossiers par horizon", ""])
    for comparison in payload["comparisons"]:
        lines.append(f"### {comparison['asset_class']} — {comparison['horizon']}")
        lines.append("| Rang | Instrument | Décision | Score | Potentiel | Révision consensus | Preuve | Priorité CI |")
        lines.append("|---:|---|---|---:|---:|---:|---|---|")
        for rank, row in enumerate(comparison["candidates"], 1):
            lines.append(
                f"| {rank} | {str(row.get('name') or row.get('isin')).replace('|', '/')} | {row.get('decision')} | {_fmt(row.get('score'))} | "
                f"{_fmt(row.get('target_upside_pct'), '%')} | {_fmt_delta(row.get('consensus_delta_4w'))} | {row.get('data_confidence')} | {row.get('action_bucket')} |"
            )
        lines.append("")

    lines.extend(["## 5. Blocages et avertissements", ""])
    issues = payload["blockers"] + payload["warnings"]
    lines.extend([f"- {item}" for item in issues] if issues else ["- Aucun blocage technique détecté ; décision finale humaine."])
    lines.extend(
        [
            "",
            "## 6. Artefacts de référence",
            "",
            "- `outputs/committee_master/COMMITTEE_DECISIONS.csv`",
            "- `outputs/committee_master/CI_COMITE_INVESTISSEMENT.docx` — analyse détaillée V2",
            "- `outputs/committee_master/CI_REFERENTIEL_PONDERE.xlsx` — référentiel pondéré complet",
            "- `outputs/decision_brief/CI_DECISION_BRIEF_V3.docx` — brief exécutif V3",
            "- `outputs/decision_brief/CI_DECISION_MATRIX_V3.csv` — matrice décisionnelle courante",
            "- `outputs/mobile/ANDROID_CI_CONTROL_CENTER.md`",
            "",
        ]
    )
    return "\n".join(lines)


def _write_word(path: Path, payload: dict) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(9)

    title = document.add_heading("Comité d’investissement — Brief décisionnel", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        f"{VERSION} — run {payload['run_id']} — {payload['selected_count']} sélection(s)"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Lecture décisionnelle uniquement : aucun score, poids, seuil ou ordre n’est modifié. "
        "La qualité de preuve qualifie la documentation disponible et ne constitue pas un score financier."
    )

    document.add_heading("1. Décisions à prendre", level=1)
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["Priorité CI", "Actif", "Horizon", "Instrument", "Décision", "Score", "Δ S/S-1", "Preuve"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    priority = {"ACTION IMMEDIATE": 0, "ATTENDRE DECLENCHEUR": 1, "CONSERVER / PROTEGER": 2, "SURVEILLER": 3, "REEXAMINER / EVITER": 4, "SORTIR / EVITER": 5, "REVUE CI": 6}
    rows = sorted(payload["decision_rows"], key=lambda row: (priority.get(str(row.get("action_bucket")), 9), -(_number(row.get("score")) or -1.0)))
    for row in rows:
        cells = table.add_row().cells
        values = [
            row.get("action_bucket"), row.get("asset_class"), row.get("horizon"), row.get("name"), row.get("decision"),
            _fmt(row.get("score")), _fmt_delta(row.get("score_delta")), row.get("data_confidence"),
        ]
        for idx, value in enumerate(values):
            cells[idx].text = _text(value, "")

    document.add_heading("2. Changements depuis S-1", level=1)
    changed = [row for row in rows if row.get("change_state") != "STABLE"]
    if not changed and not payload["removed_from_selection"]:
        document.add_paragraph("Aucun changement matériel parmi les sélections courantes.")
    for row in changed:
        document.add_paragraph(
            f"{row['name']} — {row['change_state']} — décision {row['previous_decision']} → {row['decision']} ; "
            f"score {_fmt(row.get('previous_score'))} → {_fmt(row.get('score'))} ({_fmt_delta(row.get('score_delta'))} pt).",
            style="List Bullet",
        )
    for row in payload["removed_from_selection"]:
        document.add_paragraph(
            f"{row['name']} — SORTIE_SELECTION — ancienne décision {row.get('previous_decision')} ; ancien score {_fmt(row.get('previous_score'))}.",
            style="List Bullet",
        )

    document.add_heading("3. Dossiers nécessitant une décision ou une exception", level=1)
    focus = [row for row in rows if row.get("action_bucket") in {"ACTION IMMEDIATE", "ATTENDRE DECLENCHEUR", "CONSERVER / PROTEGER", "SORTIR / EVITER"} or row.get("change_state") != "STABLE"]
    if not focus:
        document.add_paragraph("Aucun dossier d’exception ; WATCH stables résumés dans la matrice.")
    for row in focus:
        document.add_heading(f"{row['name']} — {row['asset_class']} {row['horizon']}", level=2)
        document.add_paragraph(f"Priorité CI : {row['action_bucket']} | Décision moteur : {row['decision']} | Score : {_fmt(row.get('score'))}/100")
        document.add_paragraph(f"Pourquoi maintenant : {row['why_now']}")
        document.add_paragraph(f"Facteurs moteurs : {row['positive_factors']}")
        document.add_paragraph(f"Points faibles : {row['weak_factors']}")
        document.add_paragraph(f"Invalidation / vigilance : {row['invalidation']}")
        document.add_paragraph(f"Qualité de preuve : {row['data_confidence']} — {row['data_confidence_reason']}")

    document.add_heading("4. Comparatif des meilleurs dossiers par horizon", level=1)
    for comparison in payload["comparisons"]:
        document.add_heading(f"{comparison['asset_class']} — {comparison['horizon']}", level=2)
        table = document.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["Rang", "Instrument", "Décision", "Score", "Potentiel", "Consensus 4s", "Priorité CI"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
        for rank, row in enumerate(comparison["candidates"], 1):
            cells = table.add_row().cells
            values = [
                rank, row.get("name"), row.get("decision"), _fmt(row.get("score")),
                _fmt(row.get("target_upside_pct"), "%"), _fmt_delta(row.get("consensus_delta_4w")), row.get("action_bucket"),
            ]
            for idx, value in enumerate(values):
                cells[idx].text = str(value)

    document.add_heading("5. Gouvernance", level=1)
    document.add_paragraph(
        "Cette couche V3 est un post-traitement déterministe des sorties déjà produites. Elle n’appelle aucune source externe, "
        "ne recalcule aucun score de sélection et ne peut ni créer ni modifier un ordre. Les signaux PREOPEN/POSTMARKET restent SHADOW."
    )
    document.save(path)


def _write_matrix(path: Path, rows: list[dict]) -> None:
    frame = pd.DataFrame(rows)
    if frame.empty:
        frame = pd.DataFrame(columns=["asset_class", "horizon", "isin", "name", "decision", "score", "action_bucket", "change_state"])
    frame.to_csv(path, sep=";", encoding="utf-8-sig", index=False)


def _write_snapshot(path: Path, rows: list[dict], generated_at: str) -> None:
    if not rows:
        return
    data = []
    for row in rows:
        data.append(
            {
                "generated_at_utc": generated_at,
                "asset_class": row.get("asset_class"),
                "horizon": row.get("horizon"),
                "isin": row.get("isin"),
                "name": row.get("name"),
                "decision": row.get("decision"),
                "score": row.get("score"),
                "coverage_pct": row.get("coverage_pct"),
                "consensus_score": row.get("consensus_score"),
                "consensus_delta_4w": row.get("consensus_delta_4w"),
                "target_upside_pct": row.get("target_upside_pct"),
                "entry_state": row.get("entry_state"),
                "position_state": row.get("position_state"),
                "catalyst_state": row.get("catalyst_state"),
                "data_confidence": row.get("data_confidence"),
            }
        )
    path.parent.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(data, columns=SNAPSHOT_COLUMNS).to_csv(path, sep=";", encoding="utf-8-sig", index=False)


def run(root: Path = ROOT) -> dict:
    unified = _json(root / "outputs" / "unified" / "UNIFIED_SUMMARY_LATEST.json")
    decisions = _rows(root / "outputs" / "committee_master" / "COMMITTEE_DECISIONS.csv")
    explainability = _json(root / "outputs" / "audit" / "CI_EXPLAINABILITY_AUDIT.json")
    reference = _reference(root / "outputs" / "committee_master" / "CI_REFERENTIEL_PONDERE.xlsx")
    failed, skipped = _step_health(unified)
    steps = unified.get("steps") if isinstance(unified.get("steps"), dict) else {}
    pipeline_status = str(unified.get("status") or "MISSING")
    blockers: list[str] = []
    warnings: list[str] = []
    if not unified:
        blockers.append("Résumé global absent ou illisible.")
    if pipeline_status != "SUCCESS":
        blockers.append(f"Pipeline global non réussi : {pipeline_status}.")
    if failed:
        blockers.append("Étapes en échec : " + ", ".join(failed) + ".")
    if skipped:
        blockers.append("Dépendances ignorées : " + ", ".join(skipped) + ".")
    reconstruction = explainability.get("reconstruction") if isinstance(explainability.get("reconstruction"), dict) else {}
    if explainability and not reconstruction.get("within_0_02_points", False):
        warnings.append("La reconstruction des scores publiés dépasse la tolérance de 0,02 point ou manque de données.")
    if not decisions:
        warnings.append("Aucune décision canonique lisible n’a été publiée.")
    if reference.empty:
        warnings.append("Référentiel pondéré CI indisponible : qualité de preuve et facteurs détaillés non documentés dans ce brief.")

    decision_rows = _decision_rows(decisions, reference, root)
    removed = _removed_rows(root, decision_rows)
    counts = Counter(str(row.get("decision", "MISSING")) for row in decisions)
    if blockers:
        decision_status = "BLOCKED"
    elif warnings:
        decision_status = "REVIEW_WITH_WARNINGS"
    else:
        decision_status = "READY_FOR_REVIEW"

    generated_at = datetime.now(timezone.utc).isoformat()
    payload = {
        "version": VERSION,
        "generated_at_utc": generated_at,
        "run_id": unified.get("run_id", "UNKNOWN"),
        "decision_status": decision_status,
        "pipeline_status": pipeline_status,
        "steps_success": sum(1 for value in steps.values() if isinstance(value, dict) and value.get("status") == "SUCCESS"),
        "failed_steps": failed,
        "skipped_dependencies": skipped,
        "decision_counts": dict(sorted(counts.items())),
        "selected_count": len(decision_rows),
        "changed_count": sum(1 for row in decision_rows if row.get("change_state") != "STABLE"),
        "decision_rows": decision_rows,
        "removed_from_selection": removed,
        "comparisons": _comparisons(decision_rows),
        "blockers": blockers,
        "warnings": warnings,
        "data_confidence_is_non_predictive": True,
        "preopen_postmarket_context_only": True,
        "preopen_postmarket_decision_influence": 0.0,
        "preopen_postmarket_score_influence": 0.0,
        "external_collection_calls": 0,
        "real_orders_enabled": False,
        "score_or_decision_mutation": False,
        "weight_or_threshold_changes": False,
    }

    outdir = root / "outputs" / "decision_brief"
    outdir.mkdir(parents=True, exist_ok=True)
    md_path = outdir / "DECISION_BRIEF.md"
    json_path = outdir / "DECISION_BRIEF.json"
    docx_path = outdir / "CI_DECISION_BRIEF_V3.docx"
    matrix_path = outdir / "CI_DECISION_MATRIX_V3.csv"
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    md_path.write_text(_markdown(payload), encoding="utf-8")
    _write_word(docx_path, payload)
    _write_matrix(matrix_path, decision_rows)
    _write_snapshot(root / "state" / "provenance" / "CI_DECISION_SNAPSHOT.csv", decision_rows, generated_at)
    return payload


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", default=str(ROOT))
    args = parser.parse_args()
    print(json.dumps(run(Path(args.root)), ensure_ascii=False, indent=2, default=str))


if __name__ == "__main__":
    main()
