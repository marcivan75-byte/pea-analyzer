from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import json
from typing import Iterable

import numpy as np
import pandas as pd

from v182.decision.committee_master import _pct_score, active_criteria, load_registry, resolve_field
from v182.features.etf_mt_v2081 import _criterion_scores

ROOT = Path(__file__).resolve().parents[3]
CI_VERSION = "CI_RESTITUTION_V2"
SELECTED_CODES = {"BUY_CANDIDATE", "WATCH", "REVIEW", "WATCH_NOT_TOP2"}
DECISION_PRIORITY = {"BUY_CANDIDATE": 0, "WATCH": 1, "WATCH_NOT_TOP2": 2, "REVIEW": 3}
REFERENCE_COLUMNS = [
    "asset_class",
    "horizon",
    "name",
    "isin",
    "decision",
    "final_score",
    "criterion",
    "criterion_status",
    "raw_value",
    "direction",
    "criterion_score_0_100",
    "theoretical_weight_pct",
    "effective_weight_pct",
    "weighted_contribution_points",
    "source",
    "as_of",
    "evidence_level",
    "validation_status",
]


def _read(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, sep=";", encoding="utf-8-sig", low_memory=False)


def _num(value):
    try:
        x = float(value)
    except (TypeError, ValueError):
        return None
    return x if np.isfinite(x) else None


def _clean(value, default: str = "n/a") -> str:
    if value is None:
        return default
    try:
        missing = bool(pd.isna(value))
    except (TypeError, ValueError):
        missing = False
    if missing:
        return default
    text = str(value).strip()
    return text if text and text.lower() != "nan" else default


def _source_field(resolution: str, criterion: str) -> str:
    text = str(resolution or "")
    if text.startswith(("DIRECT:", "ALIAS:")):
        return text.split(":", 1)[1]
    return criterion


def _generic_details(source: pd.DataFrame, selected: pd.DataFrame, registry: dict, asset: str, horizons: Iterable[str]) -> pd.DataFrame:
    rows: list[dict] = []
    if source.empty:
        return pd.DataFrame()
    source = source.copy()
    source["isin"] = source["isin"].astype(str)
    index_by_isin = {str(v): idx for idx, v in source["isin"].items()}
    for horizon in horizons:
        chosen = selected[(selected["asset_class"].astype(str) == asset) & (selected["horizon"].astype(str) == horizon)]
        active = active_criteria(registry, horizon)
        if chosen.empty or not active:
            continue
        raw_map: dict[str, pd.Series] = {}
        score_map: dict[str, pd.Series] = {}
        resolution_map: dict[str, str] = {}
        denom = pd.Series(0.0, index=source.index, dtype=float)
        for criterion, weight, direction in active:
            values, resolution = resolve_field(source, criterion)
            resolution_map[criterion] = resolution
            if values is None:
                raw_map[criterion] = pd.Series(np.nan, index=source.index)
                score_map[criterion] = pd.Series(np.nan, index=source.index)
                continue
            scored = _pct_score(values, direction)
            raw_map[criterion] = values
            score_map[criterion] = scored
            denom += scored.notna().astype(float) * float(weight)
        for _, decision in chosen.iterrows():
            isin = str(decision.get("isin", ""))
            idx = index_by_isin.get(isin)
            if idx is None:
                continue
            row_denom = float(denom.loc[idx]) if pd.notna(denom.loc[idx]) else 0.0
            for criterion, weight, direction in active:
                cscore = _num(score_map[criterion].loc[idx])
                raw_value = raw_map[criterion].loc[idx]
                available = cscore is not None and row_denom > 0
                eff = float(weight) / row_denom if available else 0.0
                resolution = resolution_map[criterion]
                rows.append({
                    "asset_class": asset,
                    "horizon": horizon,
                    "isin": isin,
                    "name": decision.get("name"),
                    "decision": decision.get("decision"),
                    "final_score": decision.get("score"),
                    "criterion": criterion,
                    "criterion_status": "ACTIVE" if available else "MISSING",
                    "raw_value": raw_value if pd.notna(raw_value) else None,
                    "direction": direction,
                    "resolution": resolution,
                    "source_field": _source_field(resolution, criterion),
                    "criterion_score_0_100": cscore,
                    "theoretical_weight_pct": float(weight) * 100.0,
                    "effective_weight_pct": eff * 100.0,
                    "weighted_contribution_points": cscore * eff if available else None,
                    "contribution_scope": "FINAL_SCORE",
                })
    return pd.DataFrame(rows)


def _etf_mt_details(root: Path, selected: pd.DataFrame) -> pd.DataFrame:
    chosen = selected[(selected["asset_class"].astype(str) == "ETF") & (selected["horizon"].astype(str) == "MT")]
    ranking = _read(root / "outputs" / "etf_mt_v2081" / "V20.8.2_ETF_MT_DYNAMIC_RANKING.csv")
    if chosen.empty or ranking.empty:
        return pd.DataFrame()
    cfg = json.loads((root / "config" / "V20.8_ETF_MT_HIGH_PRECISION.json").read_text(encoding="utf-8"))
    criteria_cfg = cfg["dynamic_criteria"]
    expected = list(criteria_cfg)
    weights = {name: float(spec["backtested_weight"]) for name, spec in criteria_cfg.items()}
    raw = ranking.set_index("instrument_id")[expected].apply(pd.to_numeric, errors="coerce")
    scores = _criterion_scores(raw, criteria_cfg)
    denom = pd.Series(0.0, index=raw.index, dtype=float)
    for criterion, weight in weights.items():
        denom += pd.to_numeric(scores[criterion], errors="coerce").notna().astype(float) * weight
    raw_component = float(cfg["score"]["score_raw_weight"])
    rank_component = float(cfg["score"]["cross_section_rank_weight"])
    ranked = ranking.drop_duplicates("instrument_id").set_index("instrument_id")
    rows: list[dict] = []
    for _, decision in chosen.iterrows():
        isin = str(decision.get("isin", ""))
        if isin not in raw.index or isin not in ranked.index:
            continue
        row_denom = float(denom.loc[isin]) if pd.notna(denom.loc[isin]) else 0.0
        for criterion, weight in weights.items():
            cscore = _num(scores.loc[isin, criterion])
            available = cscore is not None and row_denom > 0
            eff = weight / row_denom if available else 0.0
            rows.append({
                "asset_class": "ETF",
                "horizon": "MT",
                "isin": isin,
                "name": decision.get("name"),
                "decision": decision.get("decision"),
                "final_score": decision.get("score"),
                "criterion": criterion,
                "criterion_status": "ACTIVE" if available else "MISSING",
                "raw_value": raw.loc[isin, criterion] if pd.notna(raw.loc[isin, criterion]) else None,
                "direction": criteria_cfg[criterion].get("direction", "HIGH"),
                "resolution": f"DIRECT:{criterion}",
                "source_field": criterion,
                "criterion_score_0_100": cscore,
                "theoretical_weight_pct": weight * 100.0,
                "effective_weight_pct": eff * 100.0,
                "weighted_contribution_points": cscore * eff * raw_component if available else None,
                "contribution_scope": f"FINAL_SCORE_RAW_COMPONENT_{raw_component:.2f}",
            })
        rank_score = _num(ranked.loc[isin].get("dynamic_score_rank_pct"))
        rows.append({
            "asset_class": "ETF",
            "horizon": "MT",
            "isin": isin,
            "name": decision.get("name"),
            "decision": decision.get("decision"),
            "final_score": decision.get("score"),
            "criterion": "CROSS_SECTION_RANK_COMPONENT",
            "criterion_status": "ACTIVE" if rank_score is not None else "MISSING",
            "raw_value": rank_score,
            "direction": "HIGH",
            "resolution": "DERIVED:dynamic_score_rank_pct",
            "source_field": "dynamic_score_rank_pct",
            "criterion_score_0_100": rank_score,
            "theoretical_weight_pct": rank_component * 100.0,
            "effective_weight_pct": rank_component * 100.0,
            "weighted_contribution_points": rank_score * rank_component if rank_score is not None else None,
            "contribution_scope": "FINAL_SCORE_RANK_COMPONENT",
        })
    return pd.DataFrame(rows)


def _latest_provenance(root: Path, isins: set[str], fields: set[str]) -> dict[tuple[str, str], dict]:
    path = root / "state" / "provenance" / "OBSERVATION_PROVENANCE.csv"
    if not path.exists() or not isins or not fields:
        return {}
    usecols = ["recorded_at_utc", "isin", "field", "source", "source_url", "evidence_level", "as_of", "validation_status"]
    chunks: list[pd.DataFrame] = []
    for chunk in pd.read_csv(path, sep=";", encoding="utf-8-sig", usecols=usecols, dtype=str, chunksize=200_000, low_memory=False):
        mask = chunk["isin"].isin(isins) & chunk["field"].isin(fields)
        if mask.any():
            chunks.append(chunk.loc[mask].copy())
    if not chunks:
        return {}
    frame = pd.concat(chunks, ignore_index=True)
    frame["recorded_at_utc"] = pd.to_datetime(frame["recorded_at_utc"], errors="coerce", utc=True)
    frame = frame.sort_values("recorded_at_utc").drop_duplicates(["isin", "field"], keep="last")
    return {(str(r["isin"]), str(r["field"])): r.to_dict() for _, r in frame.iterrows()}


def _attach_provenance(root: Path, detail: pd.DataFrame) -> pd.DataFrame:
    if detail.empty:
        return detail
    lookup = _latest_provenance(root, set(detail["isin"].astype(str)), set(detail["source_field"].astype(str)))
    out = detail.copy()
    for column in ["source", "source_url", "evidence_level", "as_of", "validation_status"]:
        out[column] = [lookup.get((str(r.isin), str(r.source_field)), {}).get(column) for r in out.itertuples()]
    return out


def _join_context(root: Path, selected: pd.DataFrame) -> pd.DataFrame:
    out = selected.copy()
    entry = _read(root / "outputs" / "committee_master" / "V21_8_ENTRY_EXIT_CHALLENGER.csv")
    if not entry.empty:
        keys = ["asset_class", "horizon", "isin"]
        keep = keys + [c for c in ["v21_8_entry_state", "v21_8_position_state", "v21_8_entry_reasons", "v21_8_position_reasons"] if c in entry.columns]
        out = out.merge(entry[keep].drop_duplicates(keys, keep="last"), on=keys, how="left")
    sector = _read(root / "outputs" / "committee_master" / "COMMITTEE_SECTOR_ROTATION_V2_CONTEXT.csv")
    if not sector.empty:
        keys = ["asset_class", "horizon", "isin"]
        keep = keys + [c for c in ["valuation_warning", "correction_alert"] if c in sector.columns]
        out = out.merge(sector[keep].drop_duplicates(keys, keep="last"), on=keys, how="left")
    risk = _read(root / "outputs" / "risk" / "BETA_CORRELATION_RISK_ROWS.csv")
    if not risk.empty and "isin" in risk.columns:
        keep = ["isin"] + [c for c in ["risk_beta_reliability", "risk_metric_status", "risk_verdict"] if c in risk.columns]
        out = out.merge(risk[keep].drop_duplicates("isin", keep="last"), on="isin", how="left")
    return out


def _factor_frame(detail: pd.DataFrame, isin: str, horizon: str, positive: bool, limit: int) -> pd.DataFrame:
    subset = detail[(detail["isin"].astype(str) == isin) & (detail["horizon"].astype(str) == horizon)].copy()
    subset = subset[subset["criterion_status"].astype(str) == "ACTIVE"]
    if subset.empty:
        return subset
    subset["criterion_score_0_100"] = pd.to_numeric(subset["criterion_score_0_100"], errors="coerce")
    subset["weighted_contribution_points"] = pd.to_numeric(subset["weighted_contribution_points"], errors="coerce")
    subset = subset.dropna(subset=["criterion_score_0_100"])
    if positive:
        return subset.sort_values(
            ["weighted_contribution_points", "criterion_score_0_100"],
            ascending=[False, False],
        ).head(limit)
    return subset.sort_values(
        ["criterion_score_0_100", "weighted_contribution_points"],
        ascending=[True, True],
    ).head(limit)


def _factor_summary(detail: pd.DataFrame, isin: str, horizon: str, positive: bool) -> str:
    subset = _factor_frame(detail, isin, horizon, positive, 3)
    if subset.empty:
        return "n/a"
    return ", ".join(f"{r.criterion}={float(r.criterion_score_0_100):.0f}" for r in subset.itertuples()) or "n/a"


def _warnings(row) -> list[str]:
    warnings: list[str] = []
    for field in ["risk_verdict", "valuation_warning", "correction_alert"]:
        value = _clean(getattr(row, field, None), "")
        if value:
            warnings.append(value)
    return warnings


def _decision_conclusion(decision: str) -> str:
    mapping = {
        "BUY_CANDIDATE": (
            "Conclusion CI : préconisation positive issue du moteur de sélection. "
            "L'initiation reste conditionnée aux règles d'entrée publiées et à l'absence d'invalidation ; aucun ordre n'est automatique."
        ),
        "WATCH": (
            "Conclusion CI : dossier à conserver sous surveillance. Le niveau de conviction ne justifie pas une initiation immédiate "
            "sans amélioration des facteurs faibles ou confirmation des conditions d'entrée."
        ),
        "WATCH_NOT_TOP2": (
            "Conclusion CI : dossier recevable mais non prioritaire dans le classement courant. "
            "Il reste en surveillance et ne doit pas supplanter les sélections mieux classées sans nouveau signal."
        ),
        "REVIEW": (
            "Conclusion CI : réexamen requis avant toute décision d'allocation. "
            "Les éléments disponibles ne permettent pas de traiter le dossier comme une préconisation d'achat immédiate."
        ),
    }
    return mapping.get(decision, "Conclusion CI : décision à interpréter strictement selon le statut publié par le moteur ; aucun ordre automatique.")


def _decision_comment(row, detail: pd.DataFrame) -> str:
    isin = str(row.isin)
    horizon = str(row.horizon)
    score = _num(getattr(row, "score", None))
    coverage = _num(getattr(row, "coverage_pct", None))
    active = detail[
        (detail["isin"].astype(str) == isin)
        & (detail["horizon"].astype(str) == horizon)
        & (detail["criterion_status"].astype(str) == "ACTIVE")
    ]
    missing = detail[
        (detail["isin"].astype(str) == isin)
        & (detail["horizon"].astype(str) == horizon)
        & (detail["criterion_status"].astype(str) == "MISSING")
    ]
    score_text = f"{score:.1f}/100" if score is not None else "non disponible"
    coverage_text = f"{coverage:.1f}%" if coverage is not None else "non disponible"
    return (
        f"La décision {_clean(getattr(row, 'decision', None))} repose sur le score final publié de {score_text}, "
        f"avec une couverture de référentiel de {coverage_text}. "
        f"La décomposition retenue pour le CI comporte {len(active)} critères actifs et {len(missing)} critères explicitement manquants. "
        "Le Word synthétise les facteurs les plus contributifs, les critères les moins favorables et les alertes de contexte ; "
        "le classeur CI fournit la décomposition pondérée complète sans recalcul ni modification du score."
    )


def _android(context: pd.DataFrame, detail: pd.DataFrame) -> str:
    lines = [
        "# Comité d'investissement — Android — V21.8.1",
        "",
        "> Aucun ordre réel. T1/T2 = ACTION TCT uniquement. V21.8 = aide entrée/conservation/protection/sortie.",
        "",
    ]
    frame = context.copy()
    frame["_priority"] = frame["decision"].astype(str).map(DECISION_PRIORITY).fillna(9)
    frame["_score"] = pd.to_numeric(frame["score"], errors="coerce")
    frame = frame.sort_values(["_priority", "_score"], ascending=[True, False])
    for (asset, horizon), group in frame.groupby(["asset_class", "horizon"], sort=False):
        lines.append(f"## {asset} — {horizon}")
        for row in group.head(12).itertuples():
            score = _num(row.score)
            score_text = f"{score:.1f}" if score is not None else "n/a"
            coverage = _num(getattr(row, "coverage_pct", None))
            coverage_text = f"{coverage:.1f}%" if coverage is not None else "n/a"
            isin = str(row.isin)
            warnings = _warnings(row)
            lines.append(f"- **{row.name}** ({isin}) — {row.decision} — score {score_text} — couverture {coverage_text}")
            lines.append(f"  - +: {_factor_summary(detail, isin, str(row.horizon), True)}")
            lines.append(f"  - -: {_factor_summary(detail, isin, str(row.horizon), False)}")
            lines.append(
                f"  - V21.8: entrée={getattr(row, 'v21_8_entry_state', 'n/a') or 'n/a'} ; "
                f"position={getattr(row, 'v21_8_position_state', 'n/a') or 'n/a'}"
            )
            lines.append(f"  - warnings: {' | '.join(warnings) if warnings else 'aucun warning contextuel publié'}")
        lines.append("")
    return "\n".join(lines)


def _write_word_report(path: Path, context: pd.DataFrame, detail: pd.DataFrame) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    title = document.add_heading("Comité d'investissement — Restitution décisionnelle", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        f"{CI_VERSION} — {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} — "
        f"{len(context)} actions/ETF retenus par le moteur"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Cadre de lecture", level=1)
    document.add_paragraph(
        "Ce document commente exclusivement les décisions déjà calculées par le run. "
        "Il n'ajoute aucune donnée de marché, ne modifie aucun poids, seuil, score ou décision et n'active aucun ordre réel."
    )
    document.add_paragraph(
        "La justification complète est répartie entre ce commentaire décisionnel et le fichier Excel CI_REFERENTIEL_PONDERE.xlsx, "
        "qui reprend, pour chaque action ou ETF présenté au Comité, la totalité des critères pondérés effectivement disponibles."
    )

    frame = context.copy()
    frame["_priority"] = frame["decision"].astype(str).map(DECISION_PRIORITY).fillna(9)
    frame["_score"] = pd.to_numeric(frame["score"], errors="coerce")
    frame = frame.sort_values(["_priority", "_score"], ascending=[True, False])

    document.add_heading("Synthèse de la sélection", level=1)
    summary = (
        frame.groupby(["asset_class", "horizon", "decision"], dropna=False)
        .size()
        .reset_index(name="count")
    )
    for item in summary.itertuples():
        document.add_paragraph(
            f"{item.asset_class} — {item.horizon} — {item.decision} : {item.count}",
            style="List Bullet",
        )

    document.add_page_break()
    document.add_heading("Analyse détaillée par valeur", level=1)

    for row in frame.itertuples():
        isin = str(row.isin)
        horizon = str(row.horizon)
        name = _clean(getattr(row, "name", None), isin)
        score = _num(getattr(row, "score", None))
        coverage = _num(getattr(row, "coverage_pct", None))

        document.add_heading(f"{name} — {row.asset_class} {horizon}", level=2)
        meta = document.add_paragraph()
        meta.add_run("ISIN : ").bold = True
        meta.add_run(isin)
        meta.add_run(" | Décision : ").bold = True
        meta.add_run(_clean(getattr(row, "decision", None)))
        meta.add_run(" | Score : ").bold = True
        meta.add_run(f"{score:.1f}/100" if score is not None else "n/a")
        meta.add_run(" | Couverture : ").bold = True
        meta.add_run(f"{coverage:.1f}%" if coverage is not None else "n/a")

        document.add_paragraph(_decision_comment(row, detail))

        document.add_paragraph("Facteurs les plus contributifs :", style=None).runs[0].bold = True
        positive = _factor_frame(detail, isin, horizon, True, 5)
        if positive.empty:
            document.add_paragraph("Aucun facteur actif documenté.", style="List Bullet")
        else:
            for factor in positive.itertuples():
                contribution = _num(getattr(factor, "weighted_contribution_points", None))
                criterion_score = _num(getattr(factor, "criterion_score_0_100", None))
                source = _clean(getattr(factor, "source", None), "source non publiée")
                as_of = _clean(getattr(factor, "as_of", None), "date n/a")
                document.add_paragraph(
                    f"{factor.criterion} — score critère "
                    f"{criterion_score:.1f}/100" if criterion_score is not None else f"{factor.criterion} — score critère n/a",
                    style="List Bullet",
                )
                p = document.paragraphs[-1]
                p.add_run(
                    f" ; poids effectif {_num(getattr(factor, 'effective_weight_pct', None)) or 0.0:.2f}%"
                    f" ; contribution {contribution:.2f} pt" if contribution is not None
                    else f" ; poids effectif {_num(getattr(factor, 'effective_weight_pct', None)) or 0.0:.2f}% ; contribution n/a"
                )
                p.add_run(f" ; preuve {source}, {as_of}")

        document.add_paragraph("Critères les moins favorables / points à surveiller :", style=None).runs[0].bold = True
        weak = _factor_frame(detail, isin, horizon, False, 5)
        if weak.empty:
            document.add_paragraph("Aucun critère actif documenté.", style="List Bullet")
        else:
            for factor in weak.itertuples():
                criterion_score = _num(getattr(factor, "criterion_score_0_100", None))
                document.add_paragraph(
                    f"{factor.criterion} — score critère {criterion_score:.1f}/100"
                    if criterion_score is not None
                    else f"{factor.criterion} — score critère n/a",
                    style="List Bullet",
                )

        warnings = _warnings(row)
        document.add_paragraph("Alertes de contexte :", style=None).runs[0].bold = True
        if warnings:
            for warning in warnings:
                document.add_paragraph(warning, style="List Bullet")
        else:
            document.add_paragraph("Aucune alerte contextuelle publiée par les modules de risque/rotation.", style="List Bullet")

        entry_state = _clean(getattr(row, "v21_8_entry_state", None))
        position_state = _clean(getattr(row, "v21_8_position_state", None))
        entry_reasons = _clean(getattr(row, "v21_8_entry_reasons", None), "")
        position_reasons = _clean(getattr(row, "v21_8_position_reasons", None), "")
        document.add_paragraph(
            f"Cadre V21.8 — entrée : {entry_state} ; position : {position_state}. "
            f"{('Motifs entrée : ' + entry_reasons + '. ') if entry_reasons else ''}"
            f"{('Motifs position : ' + position_reasons + '.') if position_reasons else ''}"
        )

        notes = _clean(getattr(row, "notes", None), "")
        if notes:
            document.add_paragraph(f"Note du moteur : {notes}")

        document.add_paragraph(_decision_conclusion(str(row.decision)))
        document.add_paragraph("")

    document.add_heading("Gouvernance et traçabilité", level=1)
    document.add_paragraph(
        "Le Word et l'Excel sont générés dans la même étape post-sélection à partir de COMMITTEE_DECISIONS.csv. "
        "Le générateur échoue fermé si une valeur sélectionnée ne dispose d'aucune ligne de référentiel pondéré. "
        "La restitution ne déclenche aucune collecte externe et ne peut modifier le moteur de décision."
    )
    document.save(path)


def _excel_reference(detail: pd.DataFrame) -> pd.DataFrame:
    out = detail.copy()
    for column in REFERENCE_COLUMNS:
        if column not in out.columns:
            out[column] = None
    out["_decision_priority"] = out["decision"].astype(str).map(DECISION_PRIORITY).fillna(9)
    out["_contribution"] = pd.to_numeric(out["weighted_contribution_points"], errors="coerce")
    out = out.sort_values(
        ["_decision_priority", "asset_class", "horizon", "name", "_contribution", "criterion"],
        ascending=[True, True, True, True, False, True],
    )
    return out[REFERENCE_COLUMNS]


def _write_excel_reference(path: Path, detail: pd.DataFrame) -> None:
    frame = _excel_reference(detail)
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        frame.to_excel(writer, sheet_name="Referentiel_pondere", index=False)
        ws = writer.book["Referentiel_pondere"]
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions
        ws.sheet_view.showGridLines = False
        for column_cells in ws.columns:
            values = [str(cell.value) if cell.value is not None else "" for cell in column_cells[:200]]
            width = min(max(max((len(value) for value in values), default=0) + 2, 10), 42)
            ws.column_dimensions[column_cells[0].column_letter].width = width
        header_map = {cell.value: cell.column for cell in ws[1]}
        for name in [
            "final_score",
            "criterion_score_0_100",
            "theoretical_weight_pct",
            "effective_weight_pct",
            "weighted_contribution_points",
        ]:
            col = header_map.get(name)
            if col:
                for row_idx in range(2, ws.max_row + 1):
                    ws.cell(row=row_idx, column=col).number_format = "0.00"


def _selection_keys(frame: pd.DataFrame) -> set[tuple[str, str, str]]:
    if frame.empty:
        return set()
    return {
        (str(row.asset_class), str(row.horizon), str(row.isin))
        for row in frame[["asset_class", "horizon", "isin"]].drop_duplicates().itertuples(index=False)
    }


def _reconstruction(selected: pd.DataFrame, detail: pd.DataFrame) -> dict:
    if detail.empty:
        return {"rows": 0, "within_0_02_points": False, "max_abs_delta": None}
    grouped = (
        detail.groupby(["asset_class", "horizon", "isin"], dropna=False)["weighted_contribution_points"]
        .sum(min_count=1)
        .reset_index(name="reconstructed_score")
    )
    merged = selected.merge(grouped, on=["asset_class", "horizon", "isin"], how="left")
    merged["published"] = pd.to_numeric(merged["score"], errors="coerce")
    merged["delta"] = pd.to_numeric(merged["reconstructed_score"], errors="coerce") - merged["published"]
    valid = merged["delta"].abs().dropna()
    return {
        "rows": int(len(merged)),
        "within_0_02_points": bool((valid <= 0.02).all()) if not valid.empty else False,
        "max_abs_delta": float(valid.max()) if not valid.empty else None,
    }


def run(root: Path = ROOT) -> dict:
    decisions = _read(root / "outputs" / "committee_master" / "COMMITTEE_DECISIONS.csv")
    if decisions.empty:
        return {"status": "BLOCKED_COMMITTEE_DECISIONS_MISSING", "real_orders_enabled": False}

    selected = decisions[decisions["decision"].astype(str).isin(SELECTED_CODES)].copy()
    selected = selected[selected["asset_class"].astype(str).isin({"ACTION", "ETF"})]

    action_source = _read(root / "outputs" / "V18.2_PEA_ACTIONS_MASTER_ENRICHED.csv")
    etf_source = _read(root / "outputs" / "V18.2_PEA_ETF_MASTER_ENRICHED.csv")
    action_registry = load_registry(root / "config" / "V21_ACTIONS_REFERENCE_V21_0.json")
    etf_registry = load_registry(root / "config" / "V20_7_1_ETF_CRITERIA_REGISTRY.json")
    parts = [
        _generic_details(action_source, selected, action_registry, "ACTION", ["CT", "MT", "SHORT", "TOP_DOWN"]),
        _generic_details(etf_source, selected, etf_registry, "ETF", ["CT", "SHORT", "TOP_DOWN"]),
        _etf_mt_details(root, selected),
    ]
    detail = pd.concat([p for p in parts if not p.empty], ignore_index=True) if any(not p.empty for p in parts) else pd.DataFrame()
    detail = _attach_provenance(root, detail)
    context = _join_context(root, selected)

    mobile_dir = root / "outputs" / "mobile"
    committee_dir = root / "outputs" / "committee_master"
    audit_dir = root / "outputs" / "audit"
    mobile_dir.mkdir(parents=True, exist_ok=True)
    committee_dir.mkdir(parents=True, exist_ok=True)
    audit_dir.mkdir(parents=True, exist_ok=True)

    selected_keys = _selection_keys(selected)
    detail_keys = _selection_keys(detail)
    missing_reference = sorted(selected_keys - detail_keys)
    if missing_reference:
        blocked = {
            "status": "BLOCKED_CI_REFERENCE_INCOMPLETE",
            "version": CI_VERSION,
            "selected_rows": int(len(selected)),
            "missing_reference_keys": [list(key) for key in missing_reference],
            "score_or_decision_mutation": False,
            "weight_or_threshold_changes": False,
            "external_collection_calls": 0,
            "real_orders_enabled": False,
        }
        (audit_dir / "CI_EXPLAINABILITY_AUDIT.json").write_text(
            json.dumps(blocked, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        raise RuntimeError(f"CI_REFERENCE_INCOMPLETE:{missing_reference[:10]}")

    android_path = mobile_dir / "ANDROID_CI_CONTROL_CENTER.md"
    word_path = committee_dir / "CI_COMITE_INVESTISSEMENT.docx"
    excel_path = committee_dir / "CI_REFERENTIEL_PONDERE.xlsx"

    android_path.write_text(_android(context, detail), encoding="utf-8")
    _write_word_report(word_path, context, detail)
    _write_excel_reference(excel_path, detail)

    reconstruction = _reconstruction(selected, detail)
    payload = {
        "status": "SUCCESS",
        "version": CI_VERSION,
        "selected_rows": int(len(selected)),
        "criteria_detail_rows": int(len(detail)),
        "android_output": str(android_path.relative_to(root)),
        "word_output": str(word_path.relative_to(root)),
        "excel_output": str(excel_path.relative_to(root)),
        "same_canonical_run_android_word_excel": True,
        "same_selected_set_word_excel": selected_keys == detail_keys,
        "reference_complete_for_selected": selected_keys <= detail_keys,
        "excel_visible_sheets": ["Referentiel_pondere"],
        "score_or_decision_mutation": False,
        "weight_or_threshold_changes": False,
        "external_collection_calls": 0,
        "t1_t2_scope": "ACTION_TCT_ONLY",
        "real_orders_enabled": False,
        "reconstruction": reconstruction,
    }
    (audit_dir / "CI_EXPLAINABILITY_AUDIT.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return payload


if __name__ == "__main__":
    print(json.dumps(run(), ensure_ascii=False, indent=2))
